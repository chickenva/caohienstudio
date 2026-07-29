import math
import os
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


INPUT_DOCX = Path(r"D:\22110097_HoVuAnh\TLCN_KLTN\TLCN\CNPM_CLC_Nhom05_BaoCaoTLCN.docx")
OUT_DIR = Path(r"D:\22110097_HoVuAnh\TLCN_KLTN\TLCN\caohienstudio\output\documents")
OUTPUT_DOCX = OUT_DIR / "CNPM_CLC_Nhom05_BaoCaoTLCN_DaSua_Booking.docx"
ASSET_DIR = OUT_DIR / "generated_assets"

FIG_LABEL = "Hình"
TABLE_LABEL = "Bảng"


def load_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill="#1f2937", line_gap=6):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 28)
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, fill="#374151", width=3):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    left = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    right = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=fill)


def create_usecase_diagram(path):
    w, h = 1800, 1120
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    actor_font = load_font(27, True)
    uc_font = load_font(24)
    group_font = load_font(28, True)

    d.text((60, 36), "Lược đồ chức năng hệ thống Cao Hiển Studio", font=title_font, fill="#111827")
    d.rounded_rectangle((360, 120, 1440, 1060), radius=28, outline="#94a3b8", width=4, fill="#f8fafc")
    d.text((735, 142), "Hệ thống Website Cao Hiển Studio", font=group_font, fill="#0f172a")

    actors = {
        "Guest": (135, 280),
        "Customer": (135, 610),
        "Admin": (1665, 520),
    }

    def draw_actor(label, center):
        x, y = center
        d.ellipse((x - 28, y - 95, x + 28, y - 39), outline="#111827", width=4)
        d.line((x, y - 38, x, y + 55), fill="#111827", width=4)
        d.line((x - 55, y - 8, x + 55, y - 8), fill="#111827", width=4)
        d.line((x, y + 55, x - 52, y + 118), fill="#111827", width=4)
        d.line((x, y + 55, x + 52, y + 118), fill="#111827", width=4)
        bbox = d.textbbox((0, 0), label, font=actor_font)
        d.text((x - (bbox[2] - bbox[0]) / 2, y + 132), label, font=actor_font, fill="#111827")

    for label, center in actors.items():
        draw_actor(label, center)

    usecases = [
        ("Xem thông tin\nstudio, FAQ", (520, 275)),
        ("Tra cứu dịch vụ\nvà gallery", (820, 275)),
        ("Tư vấn AI\nChatbot", (1120, 275)),
        ("Đăng ký /\nĐăng nhập", (520, 465)),
        ("Gửi yêu cầu\nđặt lịch", (820, 465)),
        ("Chọn hình thức,\nngày + buổi", (1120, 465)),
        ("Xem hợp đồng\nbằng link/QR", (520, 655)),
        ("Xác nhận HĐ\nvà thanh toán", (820, 655)),
        ("Theo dõi đơn /\nthanh toán lại", (1120, 655)),
        ("Quản lý đơn,\ngửi hợp đồng", (590, 865)),
        ("Dời/hủy đơn,\ncập nhật trạng thái", (900, 865)),
        ("Quản lý dịch vụ,\ngallery, khách hàng", (1220, 865)),
    ]

    uc_boxes = {}
    for text, center in usecases:
        cx, cy = center
        box = (cx - 145, cy - 58, cx + 145, cy + 58)
        uc_boxes[text] = box
        d.ellipse(box, outline="#2563eb", width=3, fill="#eff6ff")
        draw_centered_text(d, box, text.replace("\n", " "), uc_font, "#1e3a8a", 4)

    lines = [
        ("Guest", 0), ("Guest", 1), ("Guest", 2), ("Guest", 3),
        ("Customer", 1), ("Customer", 2), ("Customer", 4), ("Customer", 5),
        ("Customer", 6), ("Customer", 7), ("Customer", 8),
        ("Admin", 9), ("Admin", 10), ("Admin", 11),
    ]
    ordered_boxes = list(uc_boxes.values())
    for actor_label, idx in lines:
        ax, ay = actors[actor_label]
        box = ordered_boxes[idx]
        if actor_label == "Admin":
            start = (ax - 82, ay)
            end = (box[2], (box[1] + box[3]) / 2)
        else:
            start = (ax + 82, ay)
            end = (box[0], (box[1] + box[3]) / 2)
        d.line([start, end], fill="#64748b", width=2)

    d.rounded_rectangle((420, 1000, 1380, 1046), radius=12, fill="#ecfeff", outline="#06b6d4", width=2)
    draw_centered_text(
        d,
        (430, 1000, 1370, 1046),
        "Luồng booking mới: Customer gửi yêu cầu -> Admin gửi hợp đồng -> Customer xác nhận hợp đồng -> VNPay xác nhận cọc.",
        load_font(23),
        "#155e75",
    )
    img.save(path, quality=95)


def create_sequence_diagram(path):
    w, h = 2200, 1260
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    head_font = load_font(24, True)
    msg_font = load_font(22)
    small_font = load_font(20)
    d.text((60, 38), "Sequence Diagram BookingController - Flow đặt lịch hiện tại", font=title_font, fill="#111827")

    participants = [
        ("Customer UI", 150),
        ("BookingController", 500),
        ("Booking DB", 850),
        ("Admin UI", 1190),
        ("Contract/PDF/Email", 1530),
        ("VNPay", 1880),
    ]

    top, bottom = 130, 1180
    for label, x in participants:
        d.rounded_rectangle((x - 125, top, x + 125, top + 64), radius=12, fill="#f1f5f9", outline="#475569", width=2)
        draw_centered_text(d, (x - 125, top, x + 125, top + 64), label, head_font, "#0f172a")
        d.line((x, top + 64, x, bottom), fill="#cbd5e1", width=3)

    steps = [
        (1, 150, 500, "Gửi yêu cầu đặt lịch\n(service, type, date, session)", 235),
        (2, 500, 850, "Kiểm tra user, dịch vụ,\ntrùng lịch STUDIO/OUTDOOR", 330),
        (3, 850, 500, "Lưu Booking: REQUESTED\nchưa tạo Payment", 425),
        (4, 1190, 500, "Admin chỉnh đơn /\ngửi hợp đồng", 520),
        (5, 500, 1530, "Tạo contract_token,\nPDF, QR, email", 615),
        (6, 500, 850, "Cập nhật status:\nCONTRACT_SENT", 710),
        (7, 150, 500, "Khách mở link token\nvà xác nhận hợp đồng", 805),
        (8, 500, 850, "Tạo Payment PENDING,\nstatus WAITING_PAYMENT", 900),
        (9, 500, 1880, "Trả paymentUrl,\nchuyển sang VNPay", 995),
        (10, 1880, 500, "vnpay-return + chữ ký\nSUCCESS/FAILED", 1090),
        (11, 500, 850, "SUCCESS -> Booking CONFIRMED\nFAILED -> giữ WAITING_PAYMENT", 1175),
    ]

    for no, x1, x2, text, y in steps:
        color = "#2563eb" if x1 < x2 else "#0f766e"
        arrow(d, (x1 + (35 if x1 < x2 else -35), y), (x2 - (35 if x1 < x2 else -35), y), fill=color, width=3)
        label = f"{no}. {text}"
        tx1 = min(x1, x2) + 30
        tx2 = max(x1, x2) - 30
        lines = wrap_text(d, label.replace("\n", " "), small_font, tx2 - tx1)
        for li, line in enumerate(lines[:2]):
            d.text((tx1, y - 48 + li * 24), line, font=small_font, fill="#111827")

    d.rounded_rectangle((110, 1160, 2090, 1230), radius=14, fill="#fff7ed", outline="#f97316", width=2)
    draw_centered_text(
        d,
        (120, 1160, 2080, 1230),
        "Admin dời lịch đơn CONFIRMED bằng ngày + buổi mới; backend check trùng lịch, giữ trạng thái CONFIRMED và sinh lại PDF hợp đồng.",
        msg_font,
        "#9a3412",
    )
    img.save(path, quality=95)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_para_text(paragraph, text):
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def insert_paragraph_after(paragraph, text="", style=None):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def is_caption_text(text):
    cleaned = text.strip()
    return bool(re.match(r"^(Hình|Hình|Bảng|Bảng)\s+\d", cleaned, flags=re.IGNORECASE))


def normalize_caption_style(doc):
    # Tách những đoạn đang vừa chứa ảnh vừa chứa caption, rồi chuẩn hóa style Caption.
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        has_drawing = bool(paragraph._element.xpath(".//*[local-name()='drawing']"))
        if has_drawing and is_caption_text(text):
            for node in paragraph._element.xpath(".//*[local-name()='t']"):
                node.text = ""
            new_caption = insert_paragraph_after(paragraph, text, "Caption")
            new_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if is_caption_text(text):
            paragraph.style = "Caption"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_image_before_caption(doc, caption_text, image_path):
    paragraphs = doc.paragraphs
    target_idx = None
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == caption_text:
            target_idx = i
            break
    if target_idx is None or target_idx == 0:
        raise RuntimeError(f"Không tìm thấy caption: {caption_text}")

    image_para = paragraphs[target_idx - 1]
    clear_paragraph(image_para)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(6.45))


def replace_exact_paragraphs(doc):
    replacements = {
        "-\tGiao diện khách hàng (Customer): Phát triển các chức năng xem thông tin giới thiệu, danh sách dịch vụ gói chụp, bộ sưu tập ảnh (gallery). Cho phép khách hàng đặt lịch trực tiếp, thanh toán tiền cọc và theo dõi trạng thái đơn hàng.":
            "-\tGiao diện khách hàng (Customer): Phát triển các chức năng xem thông tin giới thiệu, danh sách dịch vụ gói chụp, bộ sưu tập ảnh (gallery). Cho phép khách hàng gửi yêu cầu đặt lịch theo hình thức Studio/Ngoại cảnh, xem hợp đồng, xác nhận hợp đồng và thanh toán tiền cọc qua VNPay sau khi Admin gửi hợp đồng.",
        "-\tSử dụng Axios để thực hiện các yêu cầu HTTP đồng bộ dữ liệu theo thời gian thực (polling trạng thái thanh toán) từ frontend đến backend.":
            "-\tSử dụng Axios để thực hiện các yêu cầu HTTP đồng bộ dữ liệu đặt lịch, kiểm tra lịch bận, xác nhận hợp đồng và cập nhật kết quả thanh toán từ frontend đến backend.",
        "-\tThiết kế và triển khai các API phục vụ quản lý người dùng, dịch vụ, bộ sưu tập, xử lý nghiệp vụ đặt lịch và kiểm tra hạn sử dụng (timeout 15 phút cho đơn pending).":
            "-\tThiết kế và triển khai các API phục vụ quản lý người dùng, dịch vụ, bộ sưu tập, xử lý nghiệp vụ đặt lịch theo ngày + buổi, gửi hợp đồng, tạo thanh toán VNPay và cập nhật trạng thái đơn hàng.",
        "-\tTích hợp thanh toán trực tuyến qua cổng VNPay để xử lý tiền cọc.":
            "-\tTích hợp thanh toán trực tuyến qua cổng VNPay để xử lý tiền cọc sau khi khách hàng xác nhận hợp đồng.",
        "Đặt lịch trực tuyến (Booking): Cho phép khách hàng chọn gói dịch vụ, chọn nhiếp ảnh gia, ngày giờ và địa điểm chụp. Hệ thống có tích hợp gợi ý địa chỉ tự động và xem dự báo thời tiết trực tiếp tại màn hình đặt lịch.":
            "Đặt lịch trực tuyến (Booking): Cho phép khách hàng chọn gói dịch vụ, hình thức chụp tại Studio hoặc Ngoại cảnh, ngày chụp và buổi chụp (sáng/chiều/cả ngày). Hệ thống tự động kiểm tra lịch bận theo từng buổi để hạn chế trùng lịch.",
        "Thanh toán trực tuyến: Tích hợp cổng thanh toán VNPay, cho phép khách hàng chọn linh hoạt mức thanh toán tiền cọc (30%, 50%) hoặc thanh toán toàn bộ (100%) hóa đơn.":
            "Thanh toán trực tuyến: Tích hợp cổng thanh toán VNPay để khách hàng thanh toán tiền cọc 30% sau khi đã xem và xác nhận hợp đồng dịch vụ.",
        "Quản lý đơn hàng: Khách hàng có thể theo dõi trạng thái đơn đặt lịch, xem chi tiết, và hủy các đơn hàng đang ở trạng thái chờ thanh toán (Pending).":
            "Quản lý đơn hàng: Khách hàng có thể theo dõi trạng thái đơn đặt lịch, xem chi tiết, mở lại link thanh toán khi đơn đang chờ cọc; các trường hợp hủy/dời lịch sau khi đã xác nhận được xử lý theo chính sách của Admin.",
        "Quản lý đơn đặt lịch & Khách hàng: Xem chi tiết, cập nhật trạng thái đơn đặt lịch, đối soát thanh toán. Hỗ trợ tính năng Admin tạo đơn đặt lịch hộ cho khách hàng vãng lai hoặc khách hàng đã có trên hệ thống.":
            "Quản lý đơn đặt lịch & Khách hàng: Xem chi tiết, chỉnh đơn trước khi khách xác nhận hợp đồng, gửi/gửi lại hợp đồng, xem QR/link hợp đồng, dời lịch cho đơn đã xác nhận, cập nhật trạng thái và đối soát thanh toán.",
        "Khi người dùng chọn thanh toán, Backend hệ thống sẽ tạo ra một URL mã hóa chứa các thông tin (Mã đơn hàng, Số tiền, Thời gian, Mã bảo mật) theo tiêu chuẩn của VNPay và trả về cho Frontend. Khách hàng sẽ được chuyển hướng sang giao diện của VNPay để thực hiện giao dịch. Sau khi hoàn tất, VNPay gửi kết quả trả về hệ thống thông qua URL callback/return, hệ thống kiểm tra chữ ký bảo mật (secure hash) để đảm bảo dữ liệu không bị giả mạo trước khi cập nhật trạng thái đơn hàng.":
            "Trong luồng hiện tại, Backend chỉ tạo URL thanh toán VNPay sau khi khách hàng mở link hợp đồng và xác nhận đồng ý hợp đồng. URL này chứa mã giao dịch, số tiền cọc, thời gian hết hạn và chữ ký bảo mật theo chuẩn VNPay. Sau khi khách hoàn tất giao dịch, VNPay chuyển hướng về hệ thống; Backend xác thực secure hash trước khi cập nhật Payment thành SUCCESS và Booking thành CONFIRMED.",
        "VNPay được tích hợp để số hóa hoàn toàn luồng thanh toán của studio. Khách hàng khi đặt lịch có thể thanh toán tiền cọc linh hoạt (30%, 50%) hoặc 100% hóa đơn. Điều này giúp giảm thiểu việc khách hàng bùng lịch (no-show) và loại bỏ hoàn toàn quá trình đối soát tiền mặt thủ công, giúp Admin quản lý doanh thu minh bạch ngay trên Dashboard.":
            "VNPay được tích hợp để số hóa bước đặt cọc sau hợp đồng. Khách hàng không thanh toán ngay khi gửi yêu cầu đặt lịch; thay vào đó Admin kiểm tra/chỉnh đơn, gửi hợp đồng PDF kèm link/QR, khách xác nhận hợp đồng rồi mới thanh toán tiền cọc 30%. Điều này giúp quy trình chốt lịch rõ ràng hơn, đồng thời vẫn hỗ trợ Admin đối soát doanh thu minh bạch trên Dashboard.",
        "Đặt lịch chụp trực tuyến một cách chủ động, tiện lợi (được phép chọn ngày giờ, địa điểm và nhiếp ảnh gia yêu thích).":
            "Gửi yêu cầu đặt lịch chụp trực tuyến một cách chủ động, tiện lợi theo gói dịch vụ, hình thức chụp, ngày chụp và buổi chụp.",
        "Thanh toán tiền cọc an toàn, nhanh chóng và dễ dàng theo dõi trạng thái đơn hàng của mình.":
            "Xem hợp đồng do Admin gửi, xác nhận hợp đồng trực tuyến và thanh toán tiền cọc an toàn qua VNPay.",
        "Thực hiện luồng thao tác đặt lịch nhiều bước: điền thông tin cá nhân, chọn gói dịch vụ, chọn địa điểm (có gợi ý tự động), và xem dự báo thời tiết tại ngày chụp.":
            "Thực hiện luồng thao tác đặt lịch nhiều bước: chọn gói dịch vụ, chọn hình thức Studio/Ngoại cảnh, chọn ngày + buổi chụp, nhập địa điểm ngoại cảnh nếu cần và xem dự báo thời tiết tại ngày chụp.",
        "Thực hiện thanh toán tiền cọc trực tuyến (30%, 50% hoặc 100%) thông qua cổng VNPay.":
            "Thực hiện thanh toán tiền cọc 30% thông qua cổng VNPay sau khi đã xác nhận hợp đồng.",
        "Truy cập trang hồ sơ cá nhân để theo dõi tiến độ các đơn đặt lịch, và có quyền hủy các đơn hàng đang ở trạng thái chờ thanh toán (Pending).":
            "Truy cập trang hồ sơ cá nhân để theo dõi tiến độ các đơn đặt lịch, xem lại trạng thái hợp đồng/thanh toán và tạo lại link thanh toán khi đơn đang chờ cọc.",
        "Có khả năng thao tác cơ bản trên trình duyệt web, biết cách điền biểu mẫu trực tuyến và quen thuộc với các thao tác thanh toán điện tử (như quét mã QR Bank).":
            "Có khả năng thao tác cơ bản trên trình duyệt web, biết cách điền biểu mẫu trực tuyến, mở link/QR hợp đồng và thực hiện thanh toán điện tử qua VNPay.",
        "Theo dõi danh sách các đơn đặt lịch; cập nhật trạng thái đơn hàng (ví dụ: Từ \"Đã đặt cọc\" sang \"Hoàn thành\"); thực hiện đối soát thanh toán.":
            "Theo dõi danh sách các đơn đặt lịch; chỉnh thông tin đơn trước khi khách xác nhận hợp đồng, gửi hợp đồng, dời lịch cho đơn đã xác nhận, hủy đơn theo chính sách và cập nhật trạng thái phục vụ.",
        "Đặt lịch chụp ảnh (Booking): Khách hàng có thể thực hiện luồng đặt lịch trực tuyến, bao gồm: chọn gói dịch vụ, chọn nhiếp ảnh gia mong muốn, chọn ngày giờ và địa điểm chụp (được hỗ trợ gợi ý địa danh Việt Nam tự động và xem dự báo thời tiết).":
            "Đặt lịch chụp ảnh (Booking): Khách hàng có thể thực hiện luồng đặt lịch trực tuyến, bao gồm: chọn gói dịch vụ, chọn hình thức chụp Studio/Ngoại cảnh, chọn ngày + buổi chụp và nhập địa điểm ngoại cảnh khi cần. Hệ thống kiểm tra lịch bận tự động theo buổi.",
        "Thanh toán tiền cọc: Hệ thống hỗ trợ khách hàng thanh toán tiền cọc (30%, 50%) hoặc thanh toán toàn bộ (100%) hóa đơn thông qua việc chuyển hướng đến cổng thanh toán VNPay.":
            "Thanh toán tiền cọc: Hệ thống hỗ trợ khách hàng thanh toán tiền cọc 30% qua VNPay sau khi khách đã xác nhận hợp đồng điện tử do Admin gửi.",
        "Quản lý đơn hàng cá nhân: Khách hàng có thể truy cập trang \"Đơn của tôi\" để theo dõi lịch sử đặt lịch, xem trạng thái đơn hàng (Đã cọc, Chờ chụp, Đã hoàn thành).":
            "Quản lý đơn hàng cá nhân: Khách hàng có thể truy cập trang \"Đơn của tôi\" để theo dõi lịch sử đặt lịch, xem trạng thái đơn hàng (Đã gửi yêu cầu, Đã gửi hợp đồng, Chờ thanh toán, Đã xác nhận, Đang thực hiện, Đã hoàn thành).",
        "Hủy đơn hàng: Cho phép khách hàng chủ động hủy các đơn đặt lịch đang ở trạng thái chờ thanh toán (Pending)":
            "Hủy đơn hàng: Khách hàng chỉ có thể hủy ở các trạng thái sớm trước khi xác nhận/thanh toán; các trường hợp đã xác nhận hợp đồng hoặc đã thanh toán phải liên hệ Admin để xử lý theo chính sách dời lịch/bảo lưu/hủy.",
        "Quản lý Đơn đặt lịch & Thanh toán: Xem toàn bộ danh sách đơn đặt lịch của khách hàng, đối soát trạng thái thanh toán VNPay. Cập nhật tiến độ của đơn hàng (xác nhận lịch, hoàn tất chụp).":
            "Quản lý Đơn đặt lịch & Thanh toán: Xem toàn bộ danh sách đơn đặt lịch, chỉnh đơn ở trạng thái REQUESTED/CONTRACT_SENT, gửi hợp đồng PDF kèm QR/link, xem lại hợp đồng, dời lịch đơn CONFIRMED, cập nhật trạng thái và đối soát thanh toán VNPay.",
        "Luồng thanh toán (kết nối với VNPay) phải diễn ra mượt mà, không bị gián đoạn hay xảy ra tình trạng \"timeout\" đột ngột gây ảnh hưởng đến giao dịch tài chính của người dùng.":
            "Luồng xác nhận hợp đồng và thanh toán VNPay phải diễn ra mượt mà; nếu giao dịch thất bại hoặc link hết hạn, khách hàng có thể tạo lại link thanh toán khi đơn vẫn ở trạng thái chờ cọc.",
        "Luồng Đặt lịch chụp ảnh: Kiểm tra logic chặn đặt lịch khi thợ chụp đã bị trùng lịch (busy slots).":
            "Luồng Đặt lịch chụp ảnh: Kiểm tra logic chặn trùng lịch theo hình thức chụp, ngày chụp và buổi chụp (Studio/Outdoor, Morning/Afternoon/Full day).",
        "Luồng Quản lý Đơn hàng: Cập nhật trạng thái đơn và xử lý hủy đơn (Pending -> Canceled).":
            "Luồng Quản lý Đơn hàng: Kiểm tra chuyển trạng thái REQUESTED -> CONTRACT_SENT -> WAITING_PAYMENT -> CONFIRMED -> IN_PROGRESS -> COMPLETED, đồng thời xử lý hủy/dời lịch đúng điều kiện.",
        "Thực hiện Đặt lịch trực tuyến: Tự do chọn gói dịch vụ, thợ chụp yêu thích, ngày giờ, địa điểm và thanh toán tiền cọc an toàn qua cổng VNPay.":
            "Thực hiện Đặt lịch trực tuyến: Tự do chọn gói dịch vụ, hình thức Studio/Ngoại cảnh, ngày + buổi chụp; sau khi Admin gửi hợp đồng, khách xác nhận và thanh toán tiền cọc an toàn qua VNPay.",
        "Quản lý lịch sử đơn hàng cá nhân, chủ động hủy đơn khi đơn còn ở trạng thái chờ thanh toán.":
            "Quản lý lịch sử đơn hàng cá nhân, xem lại trạng thái hợp đồng/thanh toán và tạo lại link thanh toán khi đơn đang chờ cọc.",
        "Quản lý Đơn đặt lịch: Theo dõi toàn bộ đơn hàng của khách, thay đổi trạng thái đơn (Đã cọc -> Đã xác nhận -> Hoàn thành) để ghi nhận doanh thu.":
            "Quản lý Đơn đặt lịch: Theo dõi toàn bộ đơn hàng của khách, gửi hợp đồng, xem QR/link hợp đồng, dời lịch đơn đã xác nhận và thay đổi trạng thái đơn theo luồng nghiệp vụ mới.",
        "Tính năng Tạo đơn hộ khách: Cho phép Admin chốt đơn và thu tiền mặt trực tiếp tại quầy dành cho khách vãng lai (không qua website).":
            "Tính năng Tạo đơn hộ khách: Cho phép Admin tạo đơn REQUESTED để xử lý như luồng thông thường hoặc tạo đơn CONFIRMED khi khách đã chốt và thanh toán cọc trực tiếp tại quầy.",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in replacements:
            set_para_text(paragraph, replacements[text])


def update_tables(doc):
    table6 = doc.tables[6]
    updates = {
        8: "Gửi yêu cầu đặt lịch chụp theo ngày + buổi (Studio/Outdoor)",
        9: "Xác nhận hợp đồng và thanh toán tiền cọc trực tuyến (VNPay)",
        10: "Theo dõi tiến độ đơn đặt lịch cá nhân",
        11: "Hủy đơn sớm trước khi xác nhận hợp đồng; sau khi xác nhận/thanh toán liên hệ Admin",
        14: "Quản lý toàn bộ Đơn đặt lịch, Hợp đồng & Thanh toán",
        15: "Tạo đơn đặt lịch hộ khách hàng",
    }
    for row_idx, text in updates.items():
        table6.rows[row_idx].cells[1].text = text

    usecase_data = {
        12: {
            1: 'Customer bấm "Gửi yêu cầu đặt lịch" tại trang xác nhận đặt lịch.',
            2: "Khách hàng chọn gói, hình thức chụp, ngày + buổi và gửi yêu cầu. Hệ thống tạo đơn ở trạng thái REQUESTED để Admin kiểm tra/chỉnh đơn và gửi hợp đồng.",
            3: "Customer đã đăng nhập; gói dịch vụ còn hoạt động; khách không có đơn đang xử lý hoặc chưa hoàn tất.",
            4: "Booking được tạo ở trạng thái REQUESTED; lịch theo ngày + buổi được giữ tạm; chưa tạo Payment/VNPay ở bước này.",
            5: "1. Customer truy cập trang Đặt lịch và chọn gói dịch vụ.\n2. Customer chọn hình thức STUDIO/OUTDOOR, ngày chụp và buổi MORNING/AFTERNOON/FULL_DAY; nếu OUTDOOR thì nhập địa điểm.\n3. Frontend gọi API busy-slots để hiển thị/khóa các buổi đã bận.\n4. Customer kiểm tra lại thông tin tại trang xác nhận và gửi yêu cầu.\n5. Backend kiểm tra đơn đang xử lý, dịch vụ hợp lệ, ngày không quá khứ và trùng lịch Studio/Outdoor.\n6. Backend lưu Booking với status REQUESTED, tổng tiền và tiền cọc mặc định 30%.\n7. Giao diện thông báo gửi yêu cầu thành công, chờ Admin xử lý và gửi hợp đồng.",
            6: "Customer có thể quay lại chỉnh thông tin trước khi gửi yêu cầu.",
            7: "+ 3.1. Buổi đã bận -> Khóa lựa chọn trên giao diện.\n+ 5.1. Có đơn đang xử lý -> Báo khách hoàn tất hoặc liên hệ Admin xử lý đơn cũ.\n+ 5.2. Race condition -> Backend chặn và trả lỗi trùng lịch.",
        },
        13: {
            1: "Customer bấm đồng ý hợp đồng tại link hợp đồng do Admin gửi.",
            2: "Tạo giao dịch cọc qua VNPay và đối soát kết quả để xác nhận lịch chụp.",
            3: "Booking đang ở trạng thái CONTRACT_SENT hoặc WAITING_PAYMENT; token hợp đồng hợp lệ; đơn chưa bị hủy.",
            4: "Nếu thanh toán thành công: Payment = SUCCESS và Booking = CONFIRMED. Nếu thất bại: Booking vẫn WAITING_PAYMENT để khách có thể tạo lại link thanh toán.",
            5: "1. Customer mở link/QR hợp đồng và xem thông tin hợp đồng PDF.\n2. Customer bấm xác nhận đồng ý hợp đồng.\n3. Backend kiểm tra token, chuyển Booking sang WAITING_PAYMENT và tạo Payment PENDING với số tiền cọc 30%.\n4. Backend sinh URL VNPay và trả về Frontend.\n5. Customer thanh toán trên VNPay.\n6. VNPay redirect về hệ thống kèm tham số giao dịch.\n7. Backend xác thực secure hash, kiểm tra response code.\n8. Nếu thành công, cập nhật Payment SUCCESS và Booking CONFIRMED.",
            6: "+ 3.1. Nếu đã có Payment PENDING còn hạn, hệ thống trả lại link thanh toán cũ.\n+ 8.1. Nếu thanh toán thất bại/hủy, Payment chuyển FAILED và khách có thể tạo link mới.",
            7: "+ 7.1. Sai secure hash -> Từ chối giao dịch vì nghi ngờ giả mạo.\n+ 8.2. Đơn đã bị hủy trước đó -> Không mở lại đơn.",
        },
        14: {
            1: 'Customer mở trang "Đơn của tôi" hoặc chi tiết đơn hàng.',
            2: "Cho phép khách hàng theo dõi trạng thái đơn, xem thông tin thanh toán/hợp đồng và hủy đơn trong các trạng thái sớm được phép.",
            3: "Customer đã đăng nhập và đơn thuộc về tài khoản hiện tại.",
            4: "Khách xem được tiến độ đơn; nếu hủy hợp lệ, Booking chuyển sang CANCELED.",
            5: "1. Customer truy cập danh sách đơn hoặc chi tiết đơn.\n2. Hệ thống hiển thị trạng thái REQUESTED, CONTRACT_SENT, WAITING_PAYMENT, CONFIRMED, IN_PROGRESS, COMPLETED hoặc CANCELED.\n3. Nếu đơn WAITING_PAYMENT, Customer có thể bấm thanh toán lại để lấy link VNPay.\n4. Nếu đơn còn REQUESTED/CONTRACT_SENT/PENDING legacy, Customer có thể hủy đơn.\n5. Backend kiểm tra quyền sở hữu và trạng thái trước khi cập nhật.",
            6: "+ 3.1. Nếu link thanh toán cũ còn hạn, hệ thống trả lại link hiện có; nếu hết hạn, tạo Payment PENDING mới.\n+ 4.1. Với đơn đã CONFIRMED, Customer phải liên hệ Admin để dời lịch, bảo lưu hoặc hủy theo chính sách.",
            7: "+ 5.1. Đơn không thuộc Customer -> Từ chối truy cập.\n+ 5.2. Trạng thái không cho hủy -> Báo lỗi không thể hủy đơn.",
        },
        18: {
            1: "Admin thao tác tại trang Quản lý đơn đặt lịch.",
            2: "Quản lý toàn bộ vòng đời đơn: kiểm tra/chỉnh đơn, gửi hợp đồng, xem QR/link hợp đồng, dời lịch, hủy đơn, cập nhật tiến độ và đối soát thanh toán.",
            3: "Admin đã đăng nhập và có quyền ADMIN.",
            4: "Trạng thái Booking thay đổi đúng theo luồng nghiệp vụ; hợp đồng/PDF/QR được tạo hoặc cập nhật khi cần.",
            5: "1. Admin xem danh sách đơn và lọc theo trạng thái.\n2. Với đơn REQUESTED/CONTRACT_SENT, Admin có thể chỉnh thông tin đơn và gửi/gửi lại hợp đồng.\n3. Khi gửi hợp đồng, Backend kiểm tra trùng lịch, tạo contract_token, sinh PDF/QR và chuyển đơn sang CONTRACT_SENT.\n4. Khi khách xác nhận hợp đồng và thanh toán thành công, đơn chuyển sang CONFIRMED.\n5. Với đơn CONFIRMED, Admin có thể dời lịch, bắt đầu chụp (IN_PROGRESS), hủy theo chính sách hoặc hoàn tất đơn (COMPLETED).\n6. Khi hoàn tất, hệ thống ghi nhận phần thanh toán còn lại nếu cần để đối soát doanh thu.",
            6: "+ 2.1. Admin có thể xem lại QR/link hợp đồng ở mọi đơn đã có contract_token.\n+ 5.1. Dời lịch đơn CONFIRMED sẽ check trùng lịch, giữ trạng thái CONFIRMED và sinh lại PDF hợp đồng.",
            7: "+ 3.1. Lịch bị trùng -> Chặn gửi hợp đồng/dời lịch và yêu cầu chọn ngày/buổi khác.\n+ 5.2. Chuyển trạng thái sai luồng -> Backend từ chối bằng INVALID_TRANSITION.",
        },
        19: {
            1: 'Admin nhấn "Lưu đơn hàng" tại /admin/orders/create.',
            2: "Xử lý trường hợp khách đặt lịch trực tiếp tại studio, qua điện thoại hoặc qua kênh ngoài website.",
            3: "Admin đã đăng nhập; thông tin khách hàng, dịch vụ, ngày + buổi và địa điểm hợp lệ.",
            4: "Booking được tạo. Nếu chọn CONFIRMED, hệ thống ghi nhận Payment thủ công thành công và tạo thông tin hợp đồng tương ứng.",
            5: "1. Admin chọn khách cũ hoặc nhập thông tin khách mới.\n2. Admin chọn gói chụp, gói đi kèm, hình thức STUDIO/OUTDOOR, ngày chụp và buổi chụp.\n3. Admin nhập tổng tiền, tiền đã thu/cọc và chọn trạng thái ban đầu REQUESTED hoặc CONFIRMED.\n4. Backend kiểm tra khách hàng, dịch vụ, đơn đang xử lý và trùng lịch.\n5. Nếu khách mới, Backend tạo tài khoản khách hàng.\n6. Backend lưu Booking. Nếu status CONFIRMED, tạo Payment MANUAL SUCCESS và sinh PDF hợp đồng.\n7. Giao diện báo tạo đơn đặt hộ thành công.",
            6: "+ 3.1. Nếu chỉ cần Admin kiểm tra/chỉnh tiếp, chọn REQUESTED để đi theo luồng gửi hợp đồng thông thường.",
            7: "+ 4.1. Studio/Ekip ngoại cảnh đã có lịch trong buổi này -> Báo lỗi yêu cầu chọn ngày/buổi khác.\n+ 4.2. Khách đang có đơn chưa hoàn tất -> Chặn tạo đơn mới.",
        },
    }
    for table_idx, rows in usecase_data.items():
        table = doc.tables[table_idx]
        for row_idx, text in rows.items():
            table.rows[row_idx].cells[1].text = text


def collect_body_captions(doc):
    figure_captions = []
    table_captions = []
    in_front_matter = True
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "PHẦN 1: MỞ ĐẦU":
            in_front_matter = False
        if in_front_matter:
            continue
        if re.match(r"^(Hình|Hình)\s+\d", text, re.IGNORECASE):
            figure_captions.append(text)
        elif re.match(r"^(Bảng|Bảng)\s+\d", text, re.IGNORECASE):
            table_captions.append(text)
    return figure_captions, table_captions


def rebuild_static_lists(doc, figure_captions, table_captions, figure_pages=None, table_pages=None):
    figure_pages = figure_pages or {}
    table_pages = table_pages or {}

    paragraphs = list(doc.paragraphs)
    fig_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "DANH MỤC HÌNH ẢNH")
    table_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "DANH MỤC BẢNG")
    first_content_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "PHẦN 1: MỞ ĐẦU")

    # Xóa danh mục bảng trước để không làm lệch vùng danh mục hình.
    for p in list(doc.paragraphs)[table_heading_idx + 1:first_content_idx]:
        remove_paragraph(p)
    for p in list(doc.paragraphs)[fig_heading_idx + 1:table_heading_idx]:
        remove_paragraph(p)

    fig_heading = next(p for p in doc.paragraphs if p.text.strip() == "DANH MỤC HÌNH ẢNH")
    anchor = fig_heading
    for caption in figure_captions:
        page = figure_pages.get(caption, "...")
        anchor = insert_paragraph_after(anchor, f"{caption}\t{page}", "table of figures")

    table_heading = next(p for p in doc.paragraphs if p.text.strip() == "DANH MỤC BẢNG")
    anchor = table_heading
    for caption in table_captions:
        page = table_pages.get(caption, "...")
        anchor = insert_paragraph_after(anchor, f"{caption}\t{page}", "table of figures")


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(INPUT_DOCX, OUTPUT_DOCX)

    usecase_png = ASSET_DIR / "usecase_booking_current.png"
    sequence_png = ASSET_DIR / "sequence_booking_current.png"
    create_usecase_diagram(usecase_png)
    create_sequence_diagram(sequence_png)

    doc = Document(OUTPUT_DOCX)
    replace_exact_paragraphs(doc)
    update_tables(doc)
    normalize_caption_style(doc)
    replace_image_before_caption(doc, f"{FIG_LABEL} 2.4. Lược đồ chức năng (UseCase)", usecase_png)
    replace_image_before_caption(doc, f"{FIG_LABEL} 2.7. Lược đồ Sequence Diagram BookingController", sequence_png)

    figure_captions, table_captions = collect_body_captions(doc)
    rebuild_static_lists(doc, figure_captions, table_captions)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")
    print(f"Figures: {len(figure_captions)}")
    print(f"Tables: {len(table_captions)}")


if __name__ == "__main__":
    main()
