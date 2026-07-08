# -*- coding: utf-8 -*-
from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "documents"
ASSET_DIR = OUT_DIR / "assets_tlcn"
OUT_DOCX = OUT_DIR / "BaoCao_TLCN_CaoHienStudio_HoanChinh.docx"

PROJECT_TITLE = "XÂY DỰNG WEBSITE QUẢN LÝ VÀ ĐẶT LỊCH DỊCH VỤ CHỤP ẢNH CHO CAO HIỂN STUDIO"
STUDENT = "HỒ VŨ ANH"
STUDENT_ID = "22110097"
ADVISOR = "ThS. [Tên giáo viên hướng dẫn]"


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(text, font, max_width, draw):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill=(45, 45, 45), spacing=8):
    x1, y1, x2, y2 = box
    lines = wrap_text(text, font, x2 - x1 - 30, draw)
    line_heights = []
    total_height = 0
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        height = bbox[3] - bbox[1]
        line_heights.append(height)
        total_height += height
    total_height += spacing * max(0, len(lines) - 1)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line, height in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, font=font, fill=fill)
        y += height + spacing


def draw_box(draw, box, title, subtitle=None, fill="#FFFFFF", outline="#A8894C"):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=4)
    title_font = load_font(34, True)
    body_font = load_font(26, False)
    x1, y1, x2, y2 = box
    if subtitle:
        draw_centered_text(draw, (x1 + 14, y1 + 20, x2 - 14, y1 + 82), title, title_font, fill=(35, 35, 35), spacing=4)
        draw_centered_text(draw, (x1 + 18, y1 + 92, x2 - 18, y2 - 14), subtitle, body_font, fill=(74, 74, 74), spacing=6)
    else:
        draw_centered_text(draw, (x1 + 14, y1 + 14, x2 - 14, y2 - 14), title, title_font, fill=(35, 35, 35), spacing=6)


def arrow(draw, start, end, fill="#4F5965", width=5):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - direction * 22, ey - 12), (ex - direction * 22, ey + 12)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 12, ey - direction * 22), (ex + 12, ey - direction * 22)]
    draw.polygon(pts, fill=fill)


def create_logo():
    img = Image.new("RGBA", (700, 700), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    gold = (193, 166, 123, 255)
    dark = (42, 42, 42, 255)
    draw.ellipse((55, 55, 645, 645), outline=gold, width=22)
    draw.ellipse((120, 120, 580, 580), outline=gold, width=5)
    font_big = load_font(230, True)
    font_small = load_font(48, True)
    for text, y, font in [("CH", 210, font_big), ("STUDIO", 465, font_small)]:
        bbox = draw.textbbox((0, 0), text, font=font)
        draw.text(((700 - (bbox[2] - bbox[0])) / 2, y), text, font=font, fill=dark)
    path = ASSET_DIR / "caohien_logo.png"
    img.save(path)
    return path


def create_architecture_diagram():
    img = Image.new("RGB", (1800, 980), "#FBFAF7")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    draw.text((70, 45), "Kiến trúc tổng quan hệ thống Cao Hiển Studio", font=title_font, fill="#2F3338")
    draw_box(draw, (80, 180, 470, 400), "Người dùng", "Khách vãng lai\nKhách hàng\nQuản trị viên", "#FFFFFF")
    draw_box(draw, (650, 150, 1130, 430), "Frontend", "React, Vite, Ant Design\nRouter, Axios\nGiao diện đặt lịch và quản trị", "#FFFDF8")
    draw_box(draw, (650, 565, 1130, 830), "Backend API", "Node.js, Express\nJWT, OTP, VNPay\nBusiness logic", "#FFFFFF")
    draw_box(draw, (1320, 175, 1700, 370), "MongoDB", "Users, Services\nBookings, Payments\nGalleries, Contacts", "#F7FBFF", "#52789A")
    draw_box(draw, (1320, 500, 1700, 820), "Dịch vụ ngoài", "VNPay\nGoogle Drive\nGmail/Nodemailer\nGemini AI\nOpen-Meteo, Photon", "#FFF8F2", "#B87943")
    arrow(draw, (470, 290), (650, 290))
    arrow(draw, (890, 430), (890, 565))
    arrow(draw, (1130, 690), (1320, 285))
    arrow(draw, (1130, 690), (1320, 660))
    path = ASSET_DIR / "architecture.png"
    img.save(path)
    return path


def create_usecase_diagram():
    img = Image.new("RGB", (1800, 1100), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    actor_font = load_font(30, True)
    uc_font = load_font(26, False)
    draw.text((70, 45), "Biểu đồ use case tổng quát", font=title_font, fill="#2F3338")
    draw.rounded_rectangle((420, 140, 1380, 980), radius=32, outline="#A8894C", width=5, fill="#FFFDF8")
    draw_centered_text(draw, (420, 150, 1380, 215), "Website Cao Hiển Studio", load_font(32, True), fill=(48, 48, 48))
    actors = [("Khách vãng lai", (90, 260)), ("Khách hàng", (90, 650)), ("Quản trị viên", (1460, 455))]
    for name, (x, y) in actors:
        draw.ellipse((x + 70, y, x + 130, y + 60), outline="#4F5965", width=4)
        draw.line((x + 100, y + 60, x + 100, y + 165), fill="#4F5965", width=4)
        draw.line((x + 35, y + 98, x + 165, y + 98), fill="#4F5965", width=4)
        draw.line((x + 100, y + 165, x + 50, y + 250), fill="#4F5965", width=4)
        draw.line((x + 100, y + 165, x + 150, y + 250), fill="#4F5965", width=4)
        draw.text((x, y + 270), name, font=actor_font, fill="#2F3338")
    usecases = [
        ("Xem dịch vụ", (500, 255, 760, 335)),
        ("Xem album", (850, 255, 1110, 335)),
        ("Đăng ký / đăng nhập", (500, 410, 825, 500)),
        ("Đặt lịch và thanh toán", (860, 410, 1225, 500)),
        ("Theo dõi đơn đặt", (500, 575, 825, 665)),
        ("Quản lý đơn hàng", (860, 575, 1225, 665)),
        ("Quản lý dịch vụ", (500, 740, 825, 830)),
        ("Quản lý album", (860, 740, 1225, 830)),
        ("Xem dashboard", (690, 875, 1040, 955)),
    ]
    for text, box in usecases:
        draw.ellipse(box, outline="#BFA16A", width=4, fill="#FFFFFF")
        draw_centered_text(draw, box, text, uc_font, fill=(35, 35, 35))
    for start, end in [
        ((250, 385), (500, 295)), ((250, 385), (850, 295)),
        ((250, 775), (660, 455)), ((250, 775), (1020, 455)), ((250, 775), (660, 620)),
        ((1460, 590), (1040, 620)), ((1460, 590), (660, 785)), ((1460, 590), (1020, 785)), ((1460, 590), (865, 915)),
    ]:
        draw.line([start, end], fill="#79818A", width=3)
    path = ASSET_DIR / "usecase.png"
    img.save(path)
    return path


def create_booking_flow_diagram():
    img = Image.new("RGB", (1900, 780), "#FBFAF7")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    draw.text((70, 45), "Luồng đặt lịch và thanh toán cọc", font=title_font, fill="#2F3338")
    boxes = [
        ("Chọn dịch vụ", "Gói chính và dịch vụ bổ sung"),
        ("Chọn lịch", "Ngày, giờ, địa điểm, kiểm tra trùng lịch studio"),
        ("Xác nhận", "Tổng tiền, tỷ lệ cọc 30% - 50% - 100%"),
        ("Thanh toán VNPay", "Tạo booking PENDING và payment PENDING"),
        ("Cập nhật trạng thái", "Thành công: DEPOSITED\nThất bại/hết hạn: CANCELED"),
    ]
    x = 70
    y = 250
    w = 310
    h = 210
    for i, (title, subtitle) in enumerate(boxes):
        draw_box(draw, (x, y, x + w, y + h), title, subtitle, "#FFFFFF")
        if i < len(boxes) - 1:
            arrow(draw, (x + w, y + h / 2), (x + w + 80, y + h / 2))
        x += w + 110
    note_font = load_font(27, False)
    note = "Hệ thống giữ chỗ 15 phút cho đơn PENDING. Khi thanh toán trả về, backend xác thực chữ ký VNPay trước khi đổi trạng thái."
    draw_centered_text(draw, (170, 570, 1730, 685), note, note_font, fill=(64, 64, 64))
    path = ASSET_DIR / "booking_flow.png"
    img.save(path)
    return path


def create_erd_diagram():
    img = Image.new("RGB", (1900, 1150), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42, True)
    draw.text((70, 45), "Mô hình dữ liệu mức khái quát", font=title_font, fill="#2F3338")
    entities = {
        "User": (80, 180, 430, 420, "email\npassword_hash\nfull_name\nphone\nrole\nportfolio"),
        "Booking": (715, 150, 1185, 500, "customer_id\nservice_id\nphotographer_ids\nstart_time, end_time\nstatus\ntotal_amount"),
        "Service": (1470, 180, 1820, 430, "name\ncategory\nbase_price\nduration_hours\nfeatures\nis_active"),
        "Payment": (715, 690, 1185, 985, "booking_id\namount\npayment_method\npayment_type\ntransaction_id\nstatus"),
        "PublicGallery": (80, 690, 430, 985, "title\ncategory\ndrive_folder_id\ncoverImage\nphotographer_id\nservice_id"),
        "Contact / OTP": (1470, 690, 1820, 985, "contact messages\notp email\nTTL 300 giây"),
    }
    for name, (x1, y1, x2, y2, fields) in entities.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="#FFFDF8", outline="#A8894C", width=4)
        draw.rectangle((x1, y1, x2, y1 + 68), fill="#BFA16A")
        draw_centered_text(draw, (x1, y1 + 5, x2, y1 + 63), name, load_font(31, True), fill=(255, 255, 255))
        draw_centered_text(draw, (x1 + 14, y1 + 85, x2 - 14, y2 - 18), fields, load_font(25, False), fill=(52, 52, 52))
    for start, end in [
        ((430, 300), (715, 300)), ((1185, 300), (1470, 300)),
        ((950, 500), (950, 690)), ((430, 820), (715, 300)),
        ((1470, 820), (1185, 830)),
    ]:
        arrow(draw, start, end, fill="#5B6570", width=4)
    path = ASSET_DIR / "erd.png"
    img.save(path)
    return path


def set_cell_text(cell, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def set_paragraph_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is not None:
        p_pr.remove(p_bdr)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_page_border(section, enabled):
    sect_pr = section._sectPr
    existing = sect_pr.first_child_found_in("w:pgBorders")
    if existing is not None:
        sect_pr.remove(existing)
    if not enabled:
        return
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ["top", "left", "bottom", "right"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "18")
        tag.set(qn("w:space"), "24")
        tag.set(qn("w:color"), "8C6B2E")
        borders.append(tag)
    sect_pr.append(borders)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    existing = sect_pr.first_child_found_in("w:pgNumType")
    if existing is not None:
        sect_pr.remove(existing)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start))
    sect_pr.append(pg_num)


def configure_section(section, border=False, header=False, page_numbers=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.2)
    set_page_border(section, border)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ""
    for p in section.footer.paragraphs:
        p.text = ""
    if header:
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("BÁO CÁO TIỂU LUẬN CHUYÊN NGÀNH")
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(11)
        set_paragraph_bottom_border(p)
    if page_numbers:
        add_page_number(section.footer.paragraphs[0])
        set_page_number_start(section, 1)


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)


def set_run_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_center(doc, text, size=13, bold=False, color=None, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_blank(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def add_part(doc, text):
    add_blank(doc, 1)
    add_center(doc, text.upper(), size=16, bold=True, after=10)


def add_chapter(doc, text):
    add_blank(doc, 1)
    add_center(doc, text.upper(), size=16, bold=True, after=10)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True)
    return p


def add_caption(doc, text, kind="table"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, bold=True)
    return p


def add_table(doc, headers, rows, widths=None, font_size=12):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            cell.width = Cm(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT)
            if widths:
                cells[idx].width = Cm(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_spec_table(doc, caption, rows):
    add_caption(doc, caption)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(4.0)
        cells[1].width = Cm(12.0)
        set_cell_text(cells[0], label, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], value, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()


def add_figure(doc, image_path, caption, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption, kind="figure")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=13)


def add_cover(doc, logo_path, second=False):
    add_center(doc, "TRƯỜNG ĐẠI HỌC SƯ PHẠM KỸ THUẬT TP. HỒ CHÍ MINH", size=13, bold=True, after=2)
    add_center(doc, "KHOA CÔNG NGHỆ THÔNG TIN", size=13, bold=True, after=2)
    add_center(doc, "BỘ MÔN CÔNG NGHỆ PHẦN MỀM", size=13, bold=True, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=Cm(3.0))
    add_blank(doc, 1)
    add_center(doc, "BÁO CÁO TIỂU LUẬN CHUYÊN NGÀNH", size=16, bold=True, after=2)
    add_center(doc, "CÔNG NGHỆ PHẦN MỀM", size=16, bold=True, after=18)
    add_center(doc, PROJECT_TITLE, size=22, bold=True, color=(192, 0, 0), after=18)
    if second:
        add_center(doc, "Ngành: Công nghệ phần mềm", size=13, bold=True, after=2)
        add_center(doc, "Lớp: [Lớp sinh viên]", size=13, bold=True, after=12)
    add_center(doc, f"Sinh viên thực hiện: {STUDENT} - {STUDENT_ID}", size=14, bold=True, after=4)
    add_center(doc, f"Giảng viên hướng dẫn: {ADVISOR}", size=14, bold=True, after=4)
    add_blank(doc, 4)
    add_center(doc, "TP. HỒ CHÍ MINH, THÁNG 07 NĂM 2026", size=13, bold=True, after=0)


def add_review_form(doc, title):
    add_center(doc, title, size=16, bold=True, color=(192, 0, 0), after=10)
    rows = [
        ("Tên đề tài", PROJECT_TITLE),
        ("Sinh viên thực hiện", f"{STUDENT} - {STUDENT_ID}"),
        ("Giảng viên", "................................................................................"),
        ("Nhận xét chung", "\n".join(["........................................................................................................"] * 6)),
        ("Ưu điểm", "\n".join(["........................................................................................................"] * 4)),
        ("Hạn chế", "\n".join(["........................................................................................................"] * 4)),
        ("Điểm số", "................................................................................"),
    ]
    add_spec_table(doc, "Bảng 0.1. Thông tin nhận xét", rows)
    add_blank(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TP. Hồ Chí Minh, ngày ...... tháng ...... năm 2026")
    set_run_font(r, size=13, italic=True)
    add_blank(doc, 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Ký tên")
    set_run_font(r, size=13, bold=True)


def add_static_toc(doc):
    entries = [
        ("PHẦN 1. MỞ ĐẦU", 1, 0),
        ("1.1. Tính cấp thiết của đề tài", 1, 1),
        ("1.2. Mục tiêu của đề tài", 1, 1),
        ("1.3. Đối tượng và phạm vi nghiên cứu", 2, 1),
        ("1.4. Phương pháp thực hiện", 2, 1),
        ("1.5. Kết quả dự kiến và kết quả đạt được", 2, 1),
        ("PHẦN 2. NỘI DUNG", 3, 0),
        ("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 3, 0),
        ("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 6, 0),
        ("CHƯƠNG 3. XÂY DỰNG HỆ THỐNG", 12, 0),
        ("CHƯƠNG 4. KIỂM THỬ VÀ ĐÁNH GIÁ", 17, 0),
        ("PHẦN 3. KẾT LUẬN", 20, 0),
        ("TÀI LIỆU THAM KHẢO", 21, 0),
    ]
    add_center(doc, "MỤC LỤC", size=16, bold=True, after=12)
    for title, page, indent in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75 * indent)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(title)
        set_run_font(r, size=13, bold=(indent == 0))
        r = p.add_run(f"\t{page}")
        set_run_font(r, size=13)


def add_list_page(doc, title, entries):
    add_center(doc, title, size=16, bold=True, after=12)
    for label, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(label)
        set_run_font(r, size=13)
        r = p.add_run(f"\t{page}")
        set_run_font(r, size=13)


def add_front_matter(doc, logo):
    add_cover(doc, logo, second=False)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, border=True)
    add_cover(doc, logo, second=True)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, border=False)
    add_review_form(doc, "PHIẾU NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    doc.add_page_break()
    add_review_form(doc, "PHIẾU NHẬN XÉT CỦA GIẢNG VIÊN PHẢN BIỆN")
    doc.add_page_break()
    add_center(doc, "LỜI CẢM ƠN", size=16, bold=True, after=12)
    add_paragraph(doc, "Em xin gửi lời cảm ơn chân thành đến quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Sư phạm Kỹ thuật Thành phố Hồ Chí Minh đã truyền đạt kiến thức nền tảng về phân tích, thiết kế và xây dựng phần mềm trong suốt quá trình học tập.", first_line=True)
    add_paragraph(doc, f"Em cũng xin cảm ơn {ADVISOR} đã định hướng, góp ý và hỗ trợ em trong quá trình thực hiện đề tài. Những góp ý về nghiệp vụ, kỹ thuật và cách trình bày báo cáo là cơ sở quan trọng để em hoàn thiện sản phẩm.", first_line=True)
    add_paragraph(doc, "Do thời gian thực hiện còn hạn chế, báo cáo không tránh khỏi thiếu sót. Em kính mong nhận được sự góp ý của thầy cô để đề tài được hoàn thiện hơn trong các giai đoạn tiếp theo.", first_line=True)
    doc.add_page_break()
    add_center(doc, "ĐỀ CƯƠNG TIỂU LUẬN CHUYÊN NGÀNH", size=16, bold=True, after=12)
    add_caption(doc, "Bảng 0.2. Tóm tắt đề cương thực hiện")
    add_table(
        doc,
        ["Nội dung", "Mô tả"],
        [
            ("Tên đề tài", PROJECT_TITLE),
            ("Mục tiêu", "Xây dựng website hỗ trợ giới thiệu dịch vụ, đặt lịch, thanh toán cọc và quản trị hoạt động cho Cao Hiển Studio."),
            ("Đối tượng sử dụng", "Khách vãng lai, khách hàng đã đăng ký, quản trị viên studio và nhân sự chụp ảnh/quay phim."),
            ("Phạm vi", "Ứng dụng web gồm frontend React, backend Express, cơ sở dữ liệu MongoDB và các tích hợp VNPay, Google Drive, Gmail, Gemini AI."),
            ("Sản phẩm bàn giao", "Mã nguồn frontend/backend, cơ sở dữ liệu thiết kế, giao diện người dùng, các API nghiệp vụ và báo cáo tiểu luận chuyên ngành."),
        ],
        widths=[4.0, 12.0],
    )
    add_caption(doc, "Bảng 0.3. Kế hoạch thực hiện")
    add_table(
        doc,
        ["Giai đoạn", "Nội dung công việc", "Kết quả"],
        [
            ("Tuần 1", "Khảo sát quy trình đặt lịch và nhu cầu quản lý studio.", "Xác định bài toán và phạm vi."),
            ("Tuần 2", "Phân tích yêu cầu, tác nhân và các ca sử dụng chính.", "Danh sách chức năng và luồng nghiệp vụ."),
            ("Tuần 3", "Thiết kế cơ sở dữ liệu và kiến trúc hệ thống.", "Mô hình dữ liệu, phân rã module."),
            ("Tuần 4-5", "Xây dựng backend API, xác thực, dịch vụ, booking và thanh toán.", "Các API nghiệp vụ cốt lõi."),
            ("Tuần 6", "Xây dựng giao diện khách hàng và quản trị viên.", "Các màn hình chính của hệ thống."),
            ("Tuần 7", "Tích hợp Google Drive, email OTP, dashboard và AI chat.", "Các chức năng hỗ trợ hoàn chỉnh."),
            ("Tuần 8", "Kiểm thử, chỉnh sửa lỗi và hoàn thiện báo cáo.", "Sản phẩm và báo cáo TLCN."),
        ],
        widths=[3.0, 8.0, 5.0],
    )
    doc.add_page_break()
    add_static_toc(doc)
    doc.add_page_break()
    add_list_page(
        doc,
        "DANH MỤC HÌNH",
        [
            ("Hình 2.1. Kiến trúc tổng quan hệ thống", 7),
            ("Hình 2.2. Biểu đồ use case tổng quát", 8),
            ("Hình 2.3. Mô hình dữ liệu mức khái quát", 10),
            ("Hình 3.1. Luồng đặt lịch và thanh toán cọc", 13),
        ],
    )
    doc.add_page_break()
    add_list_page(
        doc,
        "DANH MỤC BẢNG",
        [
            ("Bảng 1.1. Công nghệ sử dụng trong hệ thống", 4),
            ("Bảng 2.1. Tác nhân của hệ thống", 6),
            ("Bảng 2.2. Yêu cầu chức năng theo nhóm người dùng", 6),
            ("Bảng 2.3. Yêu cầu phi chức năng", 7),
            ("Bảng 2.4. Đặc tả use case đăng ký tài khoản", 8),
            ("Bảng 2.5. Đặc tả use case đặt lịch", 8),
            ("Bảng 2.6. Đặc tả use case thanh toán cọc", 9),
            ("Bảng 2.7. Đặc tả use case quản lý đơn hàng", 9),
            ("Bảng 2.8. Mô tả các collection chính", 10),
            ("Bảng 2.9. Nhóm API chính của backend", 11),
            ("Bảng 3.1. Nhóm chức năng backend", 12),
            ("Bảng 3.2. Trạng thái booking", 13),
            ("Bảng 3.3. Nhóm giao diện frontend", 13),
            ("Bảng 3.4. Các tích hợp bên ngoài", 15),
            ("Bảng 3.5. Nhóm dịch vụ tiêu biểu", 15),
            ("Bảng 4.1. Kịch bản kiểm thử tiêu biểu", 17),
        ],
    )



def add_body(doc, images):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, border=False, header=True, page_numbers=True)

    add_part(doc, "PHẦN 1. MỞ ĐẦU")
    add_section_heading(doc, "1.1. Tính cấp thiết của đề tài")
    add_paragraph(doc, "Cao Hiển Studio là một đơn vị cung cấp dịch vụ chụp ảnh, quay phim và in ấn trong các sự kiện như cưới hỏi, chân dung, kỷ yếu và các buổi lễ gia đình. Với đặc thù dịch vụ theo lịch hẹn, mỗi đơn hàng thường cần nhiều thông tin như gói dịch vụ, thời gian, địa điểm, nhân sự phụ trách, chi phí đặt cọc và tình trạng thanh toán.")
    add_paragraph(doc, "Trong mô hình vận hành truyền thống, các thao tác tư vấn, báo giá, ghi nhận lịch chụp và xác nhận cọc thường được thực hiện qua điện thoại hoặc tin nhắn. Cách làm này dễ gây phân tán thông tin, trùng lịch studio, khó theo dõi trạng thái đơn hàng và mất nhiều thời gian khi khách cần xem portfolio hoặc so sánh các gói dịch vụ.")
    add_paragraph(doc, "Việc xây dựng một website quản lý và đặt lịch tập trung giúp studio giới thiệu dịch vụ chuyên nghiệp hơn, cho phép khách hàng chủ động xem album, chọn gói dịch vụ, đặt lịch trực tuyến, thanh toán cọc và theo dõi đơn đặt. Đồng thời, quản trị viên có công cụ để quản lý dịch vụ, album, khách hàng, nhiếp ảnh gia, đơn hàng và doanh thu một cách hệ thống.")

    add_section_heading(doc, "1.2. Mục tiêu của đề tài")
    add_paragraph(doc, "Mục tiêu tổng quát của đề tài là xây dựng một hệ thống web hỗ trợ quy trình kinh doanh cốt lõi của Cao Hiển Studio, từ khâu giới thiệu dịch vụ đến đặt lịch, thanh toán và quản trị sau đặt lịch.")
    add_bullets(doc, [
        "Xây dựng giao diện công khai để khách hàng xem thông tin studio, dịch vụ, bảng giá, album ảnh và liên hệ.",
        "Cho phép khách hàng đăng ký, đăng nhập, cập nhật hồ sơ, đặt lịch chụp và thanh toán cọc thông qua VNPay.",
        "Cung cấp trang quản trị cho admin để quản lý đơn hàng, khách hàng, dịch vụ, album công khai, danh mục dịch vụ/album, dữ liệu vận hành và số liệu tổng quan.",
        "Tích hợp các tiện ích hỗ trợ như OTP qua email, Google Drive cho album, dashboard doanh thu và trợ lý tư vấn AI.",
        "Thiết kế hệ thống theo hướng dễ mở rộng, tách biệt frontend/backend và lưu trữ dữ liệu bằng MongoDB.",
    ])

    add_section_heading(doc, "1.3. Đối tượng và phạm vi nghiên cứu")
    add_paragraph(doc, "Đối tượng sử dụng của hệ thống gồm khách vãng lai, khách hàng đã có tài khoản, quản trị viên studio và nhân sự chụp ảnh/quay phim được lưu trong hệ thống. Khách vãng lai chủ yếu xem thông tin và gửi liên hệ. Khách hàng có thể đặt lịch, thanh toán cọc và theo dõi đơn. Quản trị viên thao tác trên khu vực quản trị để vận hành dịch vụ.")
    add_paragraph(doc, "Phạm vi của đề tài tập trung vào ứng dụng web. Hệ thống chưa triển khai ứng dụng di động riêng, chưa xử lý toàn bộ nghiệp vụ kế toán nội bộ và chưa tự động hóa hoàn chỉnh việc tải ảnh lên Google Drive từ backend. Những nội dung này được xem là hướng phát triển sau khi hoàn thành phiên bản hiện tại.")

    add_section_heading(doc, "1.4. Phương pháp thực hiện")
    add_paragraph(doc, "Đề tài được thực hiện theo hướng phân tích nghiệp vụ trước, sau đó thiết kế kiến trúc, xây dựng từng module và kiểm thử theo kịch bản. Mã nguồn được rà soát theo hai phần chính: backend Node.js/Express chịu trách nhiệm xử lý nghiệp vụ và frontend React/Vite chịu trách nhiệm giao diện người dùng.")
    add_paragraph(doc, "Trong quá trình phân tích mã nguồn hiện tại, project có 101 tệp chính sau khi loại trừ thư mục phụ thuộc và build. Backend gồm 42 tệp với 10 controller, 11 route và 9 model chính. Frontend gồm 59 tệp với 39 trang và 9 component/layout chính. So với bản phân tích trước, hệ thống đã bổ sung module danh mục động, trang chính sách/hợp đồng, cơ chế email tự động và một số luồng quản trị mới.")

    add_section_heading(doc, "1.5. Kết quả dự kiến và kết quả đạt được")
    add_paragraph(doc, "Kết quả đạt được của đề tài là một website có đầy đủ luồng từ phía khách hàng đến phía quản trị viên. Người dùng có thể xem dịch vụ, xem album, tạo tài khoản, đặt lịch, đọc hợp đồng/chính sách và thanh toán cọc. Quản trị viên có thể quản lý dịch vụ, album, danh mục động, khách hàng, đơn hàng và theo dõi dashboard tổng quan.")
    add_paragraph(doc, "Ngoài các chức năng nghiệp vụ cơ bản, hệ thống đã tích hợp nhiều dịch vụ hỗ trợ như email OTP, VNPay, Google Drive, Gemini AI, tra cứu thời tiết và gợi ý địa chỉ. Những tích hợp này giúp đề tài có tính thực tiễn, phản ánh gần hơn quy trình vận hành của một studio dịch vụ ảnh cưới và sự kiện.")

    add_part(doc, "PHẦN 2. NỘI DUNG")
    add_chapter(doc, "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG")
    add_section_heading(doc, "1.1. Kiến trúc client-server và REST API")
    add_paragraph(doc, "Hệ thống được xây dựng theo mô hình client-server. Phần client là ứng dụng React chạy trên trình duyệt, gửi yêu cầu HTTP đến backend thông qua các API. Phần server nhận yêu cầu, kiểm tra dữ liệu, xử lý nghiệp vụ, truy vấn cơ sở dữ liệu và trả kết quả về cho client dưới dạng JSON.")
    add_paragraph(doc, "Cách tổ chức REST API giúp tách biệt rõ giao diện và nghiệp vụ. Các tài nguyên như người dùng, dịch vụ, album, đơn đặt lịch và thanh toán được biểu diễn bằng các endpoint riêng. Việc phân tách này thuận lợi cho kiểm thử, bảo trì và mở rộng hệ thống trong tương lai.")

    add_section_heading(doc, "1.2. Công nghệ frontend")
    add_paragraph(doc, "Frontend sử dụng React kết hợp Vite để xây dựng ứng dụng một trang. React giúp chia giao diện thành các component có thể tái sử dụng, còn Vite giúp quá trình phát triển nhanh hơn nhờ cơ chế dev server và build hiện đại. Thư viện React Router DOM được dùng để định tuyến giữa các trang công khai, trang khách hàng và trang quản trị.")
    add_paragraph(doc, "Giao diện quản trị và nhiều thành phần nhập liệu sử dụng Ant Design. Các thư viện hỗ trợ như Axios, Day.js và Recharts giúp xử lý gọi API, thời gian và biểu đồ dashboard. Về mặt trải nghiệm, giao diện được chia thành CustomerLayout cho người dùng bên ngoài và AdminLayout cho khu vực quản trị.")

    add_section_heading(doc, "1.3. Công nghệ backend")
    add_paragraph(doc, "Backend sử dụng Node.js và Express để xây dựng API. Các controller phụ trách xử lý nghiệp vụ, routes định nghĩa endpoint, middleware kiểm tra xác thực và models mô tả dữ liệu MongoDB thông qua Mongoose. Cấu trúc này phù hợp với hệ thống vừa và nhỏ, đồng thời giúp chia tách code theo từng nhóm chức năng.")
    add_paragraph(doc, "Hệ thống sử dụng JSON Web Token để xác thực phiên đăng nhập, bcrypt để băm mật khẩu, Nodemailer để gửi OTP qua email, Mongoose để thao tác dữ liệu, và các thư viện tích hợp VNPay, Google Drive, Gemini AI. Cách tiếp cận này giúp backend đảm nhiệm cả nghiệp vụ nội bộ lẫn kết nối dịch vụ bên ngoài.")

    add_section_heading(doc, "1.4. Cơ sở dữ liệu MongoDB")
    add_paragraph(doc, "MongoDB là cơ sở dữ liệu NoSQL lưu dữ liệu dưới dạng document. Đối với hệ thống studio, dữ liệu như gói dịch vụ, album, booking và thanh toán có cấu trúc tương đối linh hoạt, vì vậy MongoDB giúp thao tác nhanh và dễ mở rộng trường dữ liệu khi nghiệp vụ thay đổi.")
    add_paragraph(doc, "Mongoose được dùng để định nghĩa schema, ràng buộc kiểu dữ liệu, giá trị enum và quan hệ tham chiếu giữa các collection. Các collection chính gồm Users, Services, Bookings, Payments, PublicGalleries, Contacts và OTPs.")

    add_section_heading(doc, "1.5. Bảo mật, xác thực và thanh toán")
    add_paragraph(doc, "Về xác thực, hệ thống dùng JWT để xác định người dùng sau khi đăng nhập. Các route quản trị được bảo vệ bởi middleware kiểm tra token, trạng thái tài khoản và quyền ADMIN. Mật khẩu được băm trước khi lưu để hạn chế rủi ro khi dữ liệu bị lộ.")
    add_paragraph(doc, "Về thanh toán, hệ thống tích hợp VNPay theo hướng tạo URL thanh toán từ backend và xác thực chữ ký khi VNPay trả kết quả. Đơn đặt lịch ban đầu ở trạng thái PENDING, sau khi thanh toán thành công được chuyển sang DEPOSITED. Các đơn quá thời hạn giữ chỗ 15 phút sẽ bị hủy hoặc đánh dấu thanh toán hết hạn.")

    add_caption(doc, "Bảng 1.1. Công nghệ sử dụng trong hệ thống")
    add_table(
        doc,
        ["Nhóm", "Công nghệ", "Vai trò trong hệ thống"],
        [
            ("Frontend", "React, Vite, React Router DOM", "Xây dựng SPA, định tuyến và tổ chức các màn hình người dùng."),
            ("UI/UX", "Ant Design, CSS module theo trang", "Tạo form, bảng dữ liệu, modal, layout quản trị và giao diện khách hàng."),
            ("Backend", "Node.js, Express", "Xử lý API, nghiệp vụ booking, thanh toán, quản lý dữ liệu."),
            ("Database", "MongoDB, Mongoose", "Lưu trữ users, services, bookings, payments, galleries, contacts và OTP."),
            ("Xác thực", "JWT, bcrypt/bcryptjs", "Đăng nhập, phân quyền và mã hóa mật khẩu."),
            ("Email", "Nodemailer, Gmail SMTP", "Gửi OTP đăng ký, quên mật khẩu và đổi email."),
            ("Thanh toán", "VNPay, HMAC SHA512", "Tạo URL thanh toán và xác thực kết quả thanh toán cọc."),
            ("Lưu trữ ảnh", "Google Drive API", "Lấy danh sách ảnh từ thư mục Drive để hiển thị album công khai."),
            ("AI và tiện ích", "Gemini AI, Open-Meteo, Photon/Komoot", "Tư vấn khách hàng, dự báo thời tiết và gợi ý địa chỉ."),
        ],
        widths=[3.2, 5.0, 7.8],
    )

    doc.add_page_break()
    add_chapter(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    add_section_heading(doc, "2.1. Khảo sát hiện trạng")
    add_paragraph(doc, "Hoạt động của một studio chụp ảnh thường xoay quanh các nghiệp vụ: tư vấn gói dịch vụ, giới thiệu album mẫu, xác nhận lịch, nhận cọc, phân công nhân sự và theo dõi trạng thái thực hiện. Nếu các nghiệp vụ này được xử lý thủ công, dữ liệu dễ phân tán giữa tin nhắn, file bảng tính và ghi chú cá nhân.")
    add_paragraph(doc, "Từ mã nguồn hiện tại, hệ thống Cao Hiển Studio được thiết kế để giải quyết trực tiếp các vấn đề trên. Dịch vụ được chuẩn hóa thành các gói có giá và thời lượng. Lịch đặt có thời gian bắt đầu, kết thúc, địa điểm, nhân sự và trạng thái. Thanh toán được ghi nhận riêng qua collection Payments để phục vụ theo dõi doanh thu.")

    add_caption(doc, "Bảng 2.1. Tác nhân của hệ thống")
    add_table(
        doc,
        ["Tác nhân", "Mô tả", "Chức năng chính"],
        [
            ("Khách vãng lai", "Người chưa đăng nhập truy cập website.", "Xem trang chủ, dịch vụ, album, FAQ, gửi liên hệ, đăng ký tài khoản."),
            ("Khách hàng", "Người dùng đã có tài khoản CUSTOMER.", "Đặt lịch, thanh toán cọc, xem đơn của tôi, cập nhật hồ sơ, đổi mật khẩu."),
            ("Quản trị viên", "Tài khoản ADMIN được phân quyền quản trị.", "Quản lý booking, dịch vụ, gallery, danh mục, khách hàng và dashboard."),
            ("Nhiếp ảnh gia", "Nhân sự có vai trò PHOTOGRAPHER trong hệ thống.", "Được lưu hồ sơ, portfolio, chuyên môn và có thể được phân công vào booking."),
            ("Hệ thống ngoài", "VNPay, Google Drive, Gmail, Gemini, Open-Meteo, Photon.", "Thanh toán, lấy ảnh, gửi OTP, tư vấn AI, thời tiết và gợi ý địa điểm."),
        ],
        widths=[3.0, 5.5, 7.5],
    )

    add_caption(doc, "Bảng 2.2. Yêu cầu chức năng theo nhóm người dùng")
    add_table(
        doc,
        ["Nhóm", "Yêu cầu chức năng"],
        [
            ("Khách vãng lai", "Xem trang chủ, giới thiệu, dịch vụ, chi tiết dịch vụ, album, chi tiết album, FAQ và gửi form liên hệ."),
            ("Khách hàng", "Đăng ký bằng OTP, đăng nhập, quên mật khẩu, cập nhật hồ sơ, đặt lịch, xác nhận booking, thanh toán VNPay, xem lịch sử đơn và hủy đơn PENDING."),
            ("Quản trị viên", "Xem dashboard, tạo đơn hộ khách, cập nhật trạng thái đơn, quản lý dịch vụ, album, danh mục và khách hàng; dữ liệu nhân sự vẫn được lưu để phân công khi cần."),
            ("Hệ thống", "Kiểm tra token, kiểm tra quyền, giữ chỗ booking 15 phút, xác thực chữ ký VNPay, gửi OTP, lấy ảnh từ Google Drive và giới hạn tốc độ AI chat."),
        ],
        widths=[4.0, 12.0],
    )

    add_caption(doc, "Bảng 2.3. Yêu cầu phi chức năng")
    add_table(
        doc,
        ["Nhóm yêu cầu", "Mô tả"],
        [
            ("Bảo mật", "Mật khẩu được băm, route admin yêu cầu JWT và quyền ADMIN, OTP có thời hạn 300 giây."),
            ("Tính đúng đắn", "Booking phải kiểm tra thời gian, trạng thái hợp lệ, thanh toán phải xác thực chữ ký VNPay."),
            ("Khả dụng", "Giao diện phân tách rõ khu vực khách hàng và quản trị, hỗ trợ thông báo lỗi khi API không phản hồi."),
            ("Hiệu năng", "Frontend gọi API theo từng màn hình, backend lọc dữ liệu theo trạng thái và sử dụng aggregate cho dashboard."),
            ("Bảo trì", "Mã nguồn chia theo controllers, routes, models, pages và layouts, giúp dễ mở rộng module."),
            ("Mở rộng", "Có thể bổ sung PayOS, quản lý lịch nhân sự chi tiết, upload Google Drive và báo cáo doanh thu nâng cao."),
        ],
        widths=[4.0, 12.0],
    )

    add_section_heading(doc, "2.2. Kiến trúc tổng quan")
    add_paragraph(doc, "Kiến trúc của hệ thống được chia thành ba lớp chính: giao diện người dùng, backend API và cơ sở dữ liệu. Ngoài ra hệ thống còn kết nối với các dịch vụ ngoài để xử lý thanh toán, email, ảnh album, tư vấn AI và dữ liệu thời tiết.")
    add_figure(doc, images["architecture"], "Hình 2.1. Kiến trúc tổng quan hệ thống", width_cm=15.5)

    add_section_heading(doc, "2.3. Phân tích use case")
    add_paragraph(doc, "Các use case chính xoay quanh hai khu vực: khu vực công khai/khách hàng và khu vực quản trị. Người dùng công khai có thể xem thông tin, khách hàng có thể đặt lịch và thanh toán, còn quản trị viên chịu trách nhiệm vận hành dữ liệu nền của studio.")
    add_figure(doc, images["usecase"], "Hình 2.2. Biểu đồ use case tổng quát", width_cm=15.5)

    add_spec_table(doc, "Bảng 2.4. Đặc tả use case đăng ký tài khoản", [
        ("Actor", "Khách vãng lai"),
        ("Trigger", "Người dùng chọn chức năng đăng ký và nhập email."),
        ("Description", "Hệ thống gửi OTP đến email, xác thực OTP và tạo tài khoản CUSTOMER."),
        ("Pre-Conditions", "Email chưa tồn tại trong hệ thống, thông tin số điện thoại và mật khẩu hợp lệ."),
        ("Post-Conditions", "Tài khoản khách hàng được tạo, mật khẩu được băm và OTP đã dùng bị xóa."),
        ("Main Flow", "1. Người dùng nhập email để nhận OTP.\n2. Backend kiểm tra email và gửi OTP qua Nodemailer.\n3. Người dùng nhập OTP, họ tên, số điện thoại và mật khẩu.\n4. Backend kiểm tra OTP, validate dữ liệu, băm mật khẩu và tạo User."),
        ("Alternate Flow", "Nếu email đã tồn tại, hệ thống trả lỗi và yêu cầu dùng email khác."),
        ("Exception Flow", "Nếu OTP hết hạn hoặc sai, hệ thống từ chối đăng ký."),
    ])

    add_spec_table(doc, "Bảng 2.5. Đặc tả use case đặt lịch", [
        ("Actor", "Khách hàng"),
        ("Trigger", "Khách hàng chọn gói dịch vụ và nhấn đặt lịch."),
        ("Description", "Khách hàng chọn dịch vụ chính, dịch vụ bổ sung, thời gian, địa điểm và ghi chú."),
        ("Pre-Conditions", "Khách hàng đã đăng nhập, dịch vụ còn active, ngày giờ đặt nằm trong tương lai."),
        ("Post-Conditions", "Thông tin đặt lịch được chuyển sang bước xác nhận và sẵn sàng tạo thanh toán."),
        ("Main Flow", "1. Frontend tải danh sách dịch vụ.\n2. Khách hàng chọn gói chính và add-on.\n3. Hệ thống kiểm tra lịch bận của studio.\n4. Khách hàng nhập địa điểm, ghi chú và xác nhận thông tin.\n5. Frontend chuyển dữ liệu sang trang xác nhận booking."),
        ("Alternate Flow", "Nếu ngày được chọn đã có booking trùng thời gian, hệ thống yêu cầu chọn thời gian khác."),
        ("Exception Flow", "Nếu thiếu token hoặc dữ liệu không hợp lệ, backend trả lỗi và booking không được tạo."),
    ])

    add_spec_table(doc, "Bảng 2.6. Đặc tả use case thanh toán cọc", [
        ("Actor", "Khách hàng, VNPay"),
        ("Trigger", "Khách hàng xác nhận thanh toán tại trang BookingConfirm."),
        ("Description", "Backend tạo booking PENDING, payment PENDING và URL thanh toán VNPay."),
        ("Pre-Conditions", "Thông tin booking hợp lệ, khách hàng không có booking PENDING còn hiệu lực."),
        ("Post-Conditions", "Nếu thanh toán thành công, booking chuyển DEPOSITED; nếu thất bại hoặc quá hạn, booking bị hủy."),
        ("Main Flow", "1. Frontend gửi thông tin đặt lịch đến /bookings/create-vnpay.\n2. Backend tính tổng tiền, số tiền cọc và thời hạn giữ chỗ.\n3. Backend tạo URL VNPay có chữ ký HMAC SHA512.\n4. VNPay trả kết quả về backend.\n5. Backend xác thực chữ ký và cập nhật Payment/Booking."),
        ("Alternate Flow", "Nếu khách đóng trình duyệt, booking vẫn ở PENDING đến khi hết hạn 15 phút."),
        ("Exception Flow", "Nếu chữ ký không hợp lệ, backend từ chối cập nhật giao dịch."),
    ])

    add_spec_table(doc, "Bảng 2.7. Đặc tả use case quản lý đơn hàng", [
        ("Actor", "Quản trị viên"),
        ("Trigger", "Admin truy cập trang quản lý đơn hàng hoặc tạo đơn mới."),
        ("Description", "Admin xem danh sách booking, lọc trạng thái, xem chi tiết, cập nhật trạng thái và phân công nhân sự."),
        ("Pre-Conditions", "Admin đã đăng nhập, token hợp lệ và tài khoản đang active."),
        ("Post-Conditions", "Booking được cập nhật đúng theo trạng thái chuyển tiếp hợp lệ."),
        ("Main Flow", "1. Admin mở trang Orders.\n2. Frontend gọi API /bookings/admin/all.\n3. Admin xem chi tiết hoặc chọn cập nhật trạng thái.\n4. Backend kiểm tra chuyển trạng thái hợp lệ.\n5. Hệ thống lưu booking và có thể tạo payment thủ công khi hoàn tất."),
        ("Alternate Flow", "Admin có thể tạo booking hộ khách với khách hàng có sẵn hoặc thông tin khách mới."),
        ("Exception Flow", "Nếu trạng thái chuyển không hợp lệ, backend trả lỗi và không cập nhật dữ liệu."),
    ])

    add_section_heading(doc, "2.4. Thiết kế dữ liệu")
    add_paragraph(doc, "Cơ sở dữ liệu được thiết kế theo các collection xoay quanh nghiệp vụ đặt lịch. Booking là collection trung tâm, liên kết đến khách hàng, dịch vụ chính, dịch vụ bổ sung, nhân sự và các payment liên quan.")
    add_figure(doc, images["erd"], "Hình 2.3. Mô hình dữ liệu mức khái quát", width_cm=15.5)
    add_caption(doc, "Bảng 2.8. Mô tả các collection chính")
    add_table(
        doc,
        ["Collection", "Trường dữ liệu tiêu biểu", "Vai trò"],
        [
            ("Users", "email, password_hash, full_name, phone, role, portfolio, is_active", "Lưu tài khoản khách hàng, quản trị viên và nhiếp ảnh gia."),
            ("Services", "name, description, category, base_price, duration_hours, thumbnail, features, is_active", "Lưu danh mục gói chụp, quay, combo và in ấn."),
            ("Bookings", "customer_id, service_id, photographer_ids, start_time, end_time, location, total_amount, status", "Lưu đơn đặt lịch và trạng thái thực hiện."),
            ("Payments", "reference_id, amount, payment_method, payment_type, transaction_id, status, paid_at", "Theo dõi giao dịch cọc, thanh toán thủ công và trạng thái thanh toán."),
            ("PublicGalleries", "title, category, drive_folder_id, coverImage, photographer_id, service_ids, featured, is_active, order", "Quản lý album ảnh công khai, liên kết nhiều dịch vụ và lấy nguồn từ Google Drive."),
            ("Categories", "name, slug, type, description, is_active, order", "Lưu danh mục động cho dịch vụ và album, hỗ trợ lọc/hiển thị/reorder."),
            ("Contacts", "name, phone, email, message, status", "Lưu yêu cầu liên hệ từ khách hàng."),
            ("OTPs", "email, otp, createdAt", "Lưu OTP tạm thời, tự hết hạn sau 300 giây."),
        ],
        widths=[3.5, 7.2, 5.3],
    )

    add_section_heading(doc, "2.5. Thiết kế API")
    add_caption(doc, "Bảng 2.9. Nhóm API chính của backend")
    add_table(
        doc,
        ["Nhóm API", "Endpoint tiêu biểu", "Mục đích"],
        [
            ("Auth", "/api/auth/login, /register, /send-register-otp, /forgot-password, /me", "Đăng nhập, đăng ký, OTP, quên mật khẩu và lấy thông tin người dùng."),
            ("Booking", "/api/bookings/create-vnpay, /my-bookings, /vnpay-return, /admin/all, /:id/status", "Đặt lịch, thanh toán VNPay, xem đơn và quản trị đơn hàng."),
            ("Services", "/api/services, /api/services/:id, /api/services/admin", "Hiển thị dịch vụ công khai và CRUD dịch vụ cho admin."),
            ("Galleries", "/api/galleries, /api/galleries/:id, /api/galleries/admin", "Hiển thị album công khai và quản lý gallery."),
            ("Users", "/api/users/photographers, /api/users/admin/customers", "Xem nhiếp ảnh gia, quản lý khách hàng và dữ liệu nhân sự ở backend."),
            ("Categories", "/api/categories, /api/categories/admin, /api/categories/admin/reorder", "Quản lý danh mục động cho dịch vụ và album."),
            ("Contacts", "/api/contacts/send-otp, /verify-otp, /api/contacts", "Gửi OTP liên hệ, xác thực OTP và lưu yêu cầu tư vấn."),
            ("Drive", "/api/drive/folders/:folderId/images", "Lấy danh sách ảnh từ Google Drive, chuẩn hóa URL và cache ảnh."),
            ("Dashboard", "/api/dashboard/admin/overview", "Thống kê booking, khách hàng, dịch vụ, gallery và doanh thu."),
            ("AI Chat", "/api/ai-chat", "Tư vấn khách hàng bằng Gemini AI dựa trên dữ liệu dịch vụ và nhiếp ảnh gia."),
        ],
        widths=[3.2, 6.4, 6.4],
    )

    add_chapter(doc, "CHƯƠNG 3. XÂY DỰNG HỆ THỐNG")
    add_section_heading(doc, "3.1. Cấu trúc backend")
    add_paragraph(doc, "Backend được tổ chức theo các thư mục controllers, routes, models, middleware và services. File server.js cấu hình Express, CORS, JSON parser, kết nối MongoDB và mount các route chính. Middleware authMiddleware chịu trách nhiệm xác thực token và phân quyền admin hoặc photographer.")
    add_caption(doc, "Bảng 3.1. Nhóm chức năng backend")
    add_table(
        doc,
        ["Module", "Thành phần chính", "Nội dung xử lý"],
        [
            ("Xác thực", "authController, authRoutes, OTP, User", "Đăng ký, đăng nhập, OTP, quên mật khẩu, cập nhật hồ sơ và đổi mật khẩu."),
            ("Đặt lịch", "bookingController, Booking, Payment", "Tạo booking, kiểm tra trùng lịch, giữ chỗ 15 phút, thanh toán VNPay, cập nhật trạng thái."),
            ("Dịch vụ", "serviceController, Service", "CRUD dịch vụ, lọc theo category, soft delete bằng is_active và cập nhật thứ tự hiển thị."),
            ("Danh mục", "categoryController, Category", "Quản lý danh mục động cho SERVICE/GALLERY, bật/tắt và kéo thả reorder."),
            ("Album", "galleryController, PublicGallery, googleDriveService", "Quản lý gallery, trích xuất folder Drive, lấy ảnh trong thư mục, cache ảnh và cập nhật thứ tự."),
            ("Người dùng", "userController, User", "Quản lý khách hàng và nhiếp ảnh gia, portfolio, bật/tắt tài khoản."),
            ("Dashboard", "dashboardController", "Tổng hợp số lượng booking, trạng thái, doanh thu kỳ vọng và payment đã ghi nhận."),
            ("Liên hệ", "contactController, Contact, OTP", "Gửi OTP liên hệ, nhận form tư vấn và gửi email thông báo cho studio."),
            ("Tự động hóa email", "cronJobs, mailService", "Gửi email xác nhận, nhắc lịch cho khách/admin và tự chuyển CONFIRMED sang IN_PROGRESS cuối ngày."),
            ("AI chat", "aiChatController", "Tư vấn gói dịch vụ bằng Gemini, giới hạn tốc độ theo IP."),
        ],
        widths=[3.0, 5.0, 8.0],
    )

    add_section_heading(doc, "3.2. Luồng đặt lịch và thanh toán")
    add_paragraph(doc, "Luồng đặt lịch là chức năng trung tâm của hệ thống. Frontend Booking.jsx cho phép khách hàng chọn gói dịch vụ, add-on, ngày giờ, địa điểm và ghi chú. Trang BookingConfirm.jsx hiển thị thông tin xác nhận, tính tỷ lệ cọc theo thời điểm đặt và gửi yêu cầu tạo thanh toán đến backend.")
    add_paragraph(doc, "Backend kiểm tra dịch vụ còn hoạt động, ngày giờ hợp lệ, khách hàng không có đơn PENDING còn hiệu lực và studio không bị trùng lịch. Sau đó hệ thống tạo booking PENDING, payment PENDING và URL thanh toán VNPay. Khi VNPay trả kết quả, backend xác thực chữ ký rồi cập nhật trạng thái.")
    add_figure(doc, images["booking_flow"], "Hình 3.1. Luồng đặt lịch và thanh toán cọc", width_cm=15.5)
    add_caption(doc, "Bảng 3.2. Trạng thái booking")
    add_table(
        doc,
        ["Trạng thái", "Ý nghĩa", "Chuyển tiếp hợp lệ"],
        [
            ("PENDING", "Đơn vừa tạo, đang chờ thanh toán cọc.", "DEPOSITED hoặc CANCELED."),
            ("DEPOSITED", "Khách đã thanh toán cọc thành công.", "CONFIRMED hoặc CANCELED."),
            ("CONFIRMED", "Studio đã xác nhận lịch.", "IN_PROGRESS hoặc CANCELED."),
            ("IN_PROGRESS", "Dịch vụ đang được thực hiện.", "COMPLETED."),
            ("COMPLETED", "Đơn đã hoàn tất.", "Trạng thái kết thúc."),
            ("CANCELED", "Đơn đã bị hủy.", "Trạng thái kết thúc."),
        ],
        widths=[3.2, 7.0, 5.8],
    )

    add_section_heading(doc, "3.3. Cấu trúc frontend")
    add_paragraph(doc, "Frontend được chia thành nhóm trang xác thực, nhóm trang khách hàng và nhóm trang quản trị. App.jsx định nghĩa các route chính. CustomerLayout chứa header công khai, menu điều hướng, thông tin người dùng và AIChatWidget. AdminLayout chứa sidebar quản trị và kiểm tra quyền ADMIN trước khi cho truy cập.")
    add_caption(doc, "Bảng 3.3. Nhóm giao diện frontend")
    add_table(
        doc,
        ["Nhóm giao diện", "Trang/Component", "Chức năng"],
        [
            ("Xác thực", "Login, Register, ForgotPassword", "Đăng nhập, đăng ký bằng OTP và đặt lại mật khẩu."),
            ("Khách hàng", "Home, About, Services, ServiceDetail, Galleries, GalleryDetail, Booking, BookingConfirm, MyBookings, BookingDetail, Profile, Contact, FAQ", "Xem thông tin, đặt lịch, theo dõi booking và cập nhật tài khoản."),
            ("Chính sách", "Contract, RefundPolicy", "Hiển thị hợp đồng dịch vụ và chính sách hủy/hoàn cọc để khách xác nhận trước khi đặt lịch."),
            ("Quản trị", "AdminDashboard, AdminOrders, CreateOrder, AdminServices, ServiceForm, AdminGalleries, GalleryForm, AdminCategories, AdminCustomers, AdminProfile", "Quản lý dữ liệu vận hành studio theo route hiện tại."),
            ("Layout", "CustomerLayout, AdminLayout, Shared/Public/Landing layouts", "Tổ chức giao diện, menu, phân quyền và vùng hiển thị nội dung."),
            ("Tiện ích", "AIChatWidget, Logo", "Tư vấn AI và nhận diện thương hiệu Cao Hiển Studio."),
        ],
        widths=[3.0, 6.0, 7.0],
    )

    add_section_heading(doc, "3.4. Chức năng quản trị")
    add_paragraph(doc, "Khu vực quản trị cho phép admin theo dõi dashboard, lọc và cập nhật booking, tạo đơn hộ khách, quản lý dịch vụ, album, danh mục và khách hàng. Các thao tác quan trọng đều gọi API có middleware verifyAdmin để đảm bảo chỉ tài khoản ADMIN mới được phép thực hiện.")
    add_paragraph(doc, "Trang AdminOrders hiển thị danh sách đơn, trạng thái, thông tin khách hàng, dịch vụ, thời gian và tổng tiền. Admin có thể cập nhật trạng thái theo quy tắc chuyển tiếp được backend kiểm soát. Trang CreateOrder hỗ trợ tạo booking thủ công, phù hợp với trường hợp khách đặt trực tiếp tại studio hoặc qua kênh ngoài website.")
    add_paragraph(doc, "Trang AdminServices và ServiceForm hỗ trợ thêm, sửa, ẩn/hiện dịch vụ. Dịch vụ được phân loại thành TRADITIONAL, PHOTOJOURNALISM, COMBO, PRINT và OTHER. Trang AdminGalleries quản lý album công khai bằng Google Drive folder ID, cho phép chọn album nổi bật và bật/tắt trạng thái hiển thị.")
    add_paragraph(doc, "Phiên bản hiện tại bổ sung trang AdminCategories cho hai loại danh mục SERVICE và GALLERY. Admin có thể thêm, sửa, xóa, bật/tắt và kéo thả để cập nhật thứ tự hiển thị danh mục. Frontend Services và Galleries gọi API categories để hiển thị bộ lọc động thay vì phụ thuộc hoàn toàn vào danh mục hardcode.")

    add_section_heading(doc, "3.5. Tích hợp dịch vụ ngoài")
    add_caption(doc, "Bảng 3.4. Các tích hợp bên ngoài")
    add_table(
        doc,
        ["Tích hợp", "Vị trí sử dụng", "Mục đích"],
        [
            ("VNPay", "bookingController, VnpayReturn", "Thanh toán cọc trực tuyến, xác thực giao dịch và cập nhật trạng thái booking."),
            ("Gmail/Nodemailer", "authController, contactController", "Gửi OTP đăng ký, quên mật khẩu, cập nhật email và thông báo liên hệ."),
            ("Google Drive API", "googleDriveService, galleryController, driveController", "Lấy ảnh từ thư mục Drive để hiển thị gallery công khai."),
            ("Gemini AI", "aiChatController, AIChatWidget", "Tư vấn gói chụp, phong cách, lịch trình và gợi ý cho khách hàng."),
            ("Open-Meteo", "Booking.jsx", "Hiển thị thông tin thời tiết hỗ trợ chọn ngày chụp."),
            ("Photon/Komoot", "Booking.jsx", "Gợi ý địa chỉ khi khách nhập địa điểm chụp."),
        ],
        widths=[3.5, 5.0, 7.5],
    )

    add_section_heading(doc, "3.6. Một số dữ liệu nghiệp vụ mẫu")
    add_paragraph(doc, "Script seedRealServices.js cho thấy hệ thống đã xây dựng danh mục dịch vụ thực tế gồm chụp truyền thống, quay truyền thống, chụp/quay phóng sự, combo và in ấn. Các gói có giá, thời lượng, ảnh đại diện, mô tả và danh sách quyền lợi. Đây là dữ liệu nền giúp trang dịch vụ và luồng booking vận hành giống nhu cầu thật của studio.")
    add_caption(doc, "Bảng 3.5. Nhóm dịch vụ tiêu biểu")
    add_table(
        doc,
        ["Nhóm dịch vụ", "Ví dụ", "Ý nghĩa"],
        [
            ("TRADITIONAL", "Chụp truyền thống, quay truyền thống, gói lễ công cô, flycam", "Dịch vụ sự kiện theo phong cách truyền thống."),
            ("PHOTOJOURNALISM", "Chụp phóng sự Basic/VIP, quay phóng sự Basic/VIP", "Dịch vụ ghi lại khoảnh khắc tự nhiên và câu chuyện sự kiện."),
            ("COMBO", "Chụp phóng sự + truyền thống Basic/VIP", "Gói kết hợp nhiều máy chụp để bao phủ tốt hơn."),
            ("PRINT", "In ảnh lẻ, photobook, hình lớn ép gỗ/mika", "Dịch vụ in ấn bổ sung sau chụp."),
        ],
        widths=[4.0, 6.0, 6.0],
    )

    doc.add_page_break()
    add_section_heading(doc, "3.7. Tự động hóa email và chính sách dịch vụ")
    add_paragraph(doc, "Hệ thống đã bổ sung mailService và cronJobs để tự động hóa một phần quy trình chăm sóc khách hàng. Khi booking thanh toán thành công, hệ thống gửi email xác nhận cho khách và email báo đơn mới cho admin. Hằng ngày lúc 7:00, cron job gửi nhắc lịch cho khách có lịch chụp vào ngày mai và nhắc admin về lịch chụp trong ngày. Lúc 23:55, hệ thống tự động chuyển các booking CONFIRMED của ngày hiện tại sang IN_PROGRESS và gửi thông báo cho admin.")
    add_paragraph(doc, "Frontend cũng có hai trang văn bản pháp lý gồm hợp đồng dịch vụ và chính sách hủy/hoàn cọc. Trong luồng Booking, khách cần xác nhận đã đọc và đồng ý với các chính sách này trước khi tiếp tục. Đây là điểm bổ sung quan trọng vì quy trình đặt lịch dịch vụ ảnh cưới cần làm rõ trách nhiệm, tiền cọc, bảo lưu lịch và thanh toán phần còn lại.")

    doc.add_page_break()
    add_chapter(doc, "CHƯƠNG 4. KIỂM THỬ VÀ ĐÁNH GIÁ")
    add_section_heading(doc, "4.1. Phương pháp kiểm thử")
    add_paragraph(doc, "Đề tài được đánh giá theo hướng kiểm thử chức năng và kiểm thử luồng nghiệp vụ chính. Do project hiện chưa có bộ kiểm thử tự động hoàn chỉnh, các kịch bản trong báo cáo tập trung vào kiểm thử thủ công dựa trên API và giao diện người dùng.")
    add_paragraph(doc, "Các nhóm kiểm thử chính gồm xác thực tài khoản, đặt lịch, thanh toán, quản trị dịch vụ, quản trị album, quản trị khách hàng, dashboard và AI chat. Với mỗi kịch bản, hệ thống được đánh giá dựa trên đầu vào, hành động, kết quả mong đợi và trạng thái xử lý.")

    add_caption(doc, "Bảng 4.1. Kịch bản kiểm thử tiêu biểu")
    add_table(
        doc,
        ["Mã", "Chức năng", "Kịch bản", "Kết quả mong đợi"],
        [
            ("TC01", "Đăng ký", "Nhập email mới, nhận OTP, nhập thông tin hợp lệ.", "Tạo tài khoản CUSTOMER, OTP bị xóa sau khi dùng."),
            ("TC02", "Đăng ký", "Nhập email đã tồn tại.", "Hệ thống trả thông báo email đã được sử dụng."),
            ("TC03", "Đăng nhập", "Nhập đúng email và mật khẩu.", "Trả JWT và thông tin người dùng, điều hướng theo role."),
            ("TC04", "Quên mật khẩu", "Gửi OTP và đặt mật khẩu mới hợp lệ.", "Mật khẩu mới được băm và lưu."),
            ("TC05", "Xem dịch vụ", "Mở trang Services.", "Hiển thị danh sách dịch vụ active, lọc theo category."),
            ("TC06", "Đặt lịch", "Chọn ngày trong quá khứ.", "Backend từ chối vì thời gian không hợp lệ."),
            ("TC07", "Đặt lịch", "Chọn khung giờ đã có booking trùng.", "Hệ thống yêu cầu chọn thời gian khác."),
            ("TC08", "Thanh toán", "Tạo booking và thanh toán VNPay thành công.", "Payment SUCCESS, booking DEPOSITED."),
            ("TC09", "Thanh toán", "Booking PENDING quá 15 phút.", "Booking bị hủy hoặc payment chuyển hết hạn."),
            ("TC10", "Đơn của tôi", "Khách hàng xem danh sách booking.", "Chỉ hiển thị booking thuộc customer đang đăng nhập."),
            ("TC11", "Admin orders", "Admin cập nhật PENDING sang CONFIRMED trực tiếp.", "Backend từ chối do sai thứ tự chuyển trạng thái."),
            ("TC12", "Admin services", "Admin tạo dịch vụ mới hợp lệ.", "Dịch vụ được lưu và xuất hiện trong danh sách admin."),
            ("TC13", "Admin galleries", "Admin nhập link Google Drive folder.", "Backend trích xuất folder ID và lưu gallery."),
            ("TC14", "Dashboard", "Admin mở dashboard.", "Hiển thị số lượng booking, khách hàng, dịch vụ, gallery và doanh thu."),
            ("TC15", "AI chat", "Gửi câu hỏi tư vấn dưới 1000 ký tự.", "AI trả lời dựa trên dữ liệu dịch vụ và photographer active."),
        ],
        widths=[1.6, 3.0, 6.2, 5.2],
        font_size=11,
    )

    add_section_heading(doc, "4.2. Đánh giá kết quả")
    add_paragraph(doc, "Hệ thống đã đáp ứng được các nghiệp vụ quan trọng của một website studio: giới thiệu thương hiệu, quản lý dịch vụ, quản lý album, đăng ký/đăng nhập, đặt lịch, thanh toán cọc và quản trị booking. Cấu trúc code tách biệt frontend/backend giúp việc bảo trì thuận lợi.")
    add_paragraph(doc, "Luồng booking được xây dựng khá đầy đủ với kiểm tra thời gian trong tương lai, chặn booking PENDING còn hiệu lực, giữ chỗ 15 phút, xác thực kết quả VNPay và quản lý trạng thái đơn theo quy tắc. Dashboard admin cũng giúp chủ studio theo dõi nhanh tình hình vận hành.")

    add_section_heading(doc, "4.3. Hạn chế hiện tại")
    add_paragraph(doc, "Qua quá trình rà soát mã nguồn, hệ thống vẫn còn một số điểm cần hoàn thiện trước khi triển khai chính thức:")
    add_bullets(doc, [
        "Một số file frontend còn hardcode API URL là http://localhost:5000/api; nên chuyển hoàn toàn sang biến môi trường để dễ deploy.",
        "Frontend có logic mã giảm giá trong Booking/BookingConfirm, nhưng backend hiện chưa tính discount_amount hoặc coupon_code khi tạo thanh toán, có thể làm lệch số tiền hiển thị và số tiền backend tính.",
        "driveController có route tạo folder và upload ảnh, nhưng googleDriveService hiện mới triển khai listImagesInFolder với quyền drive.readonly; cần bổ sung scope và hàm upload nếu muốn quản trị ảnh trực tiếp.",
        "Một số file quản lý photographer vẫn còn trong frontend/backend nhưng route admin hiện tại chưa mở lại trong App.jsx; cần quyết định bật lại chức năng hoặc dọn code cũ để tránh lệch tài liệu.",
        "Một số nhãn trạng thái như EXPIRED hoặc PAYMENT_FAILED xuất hiện trong thống kê nhưng chưa nằm trong enum Booking hiện tại; cần thống nhất mô hình trạng thái.",
        "Booking model vẫn còn resource_ids tham chiếu Resource trong khi module Resource/Rentals đã được gỡ khỏi route chính; cần dọn schema hoặc khôi phục module nếu nghiệp vụ quay lại.",
        "Trang hợp đồng/chính sách mô tả cọc cố định 30%, trong khi luồng Booking/Backend hỗ trợ deposit_percent 30%, 50% hoặc 100%; cần thống nhất chính sách hiển thị và logic thanh toán.",
        "Contact API có endpoint gửi/xác thực OTP nhưng submitContact chưa tự kiểm tra OTP trong cùng request; frontend cần đảm bảo flow verify trước submit hoặc backend cần enforce lại.",
        "Project chưa có bộ unit test/integration test tự động, vì vậy nên bổ sung test cho auth, booking, payment callback và các route admin.",
    ])

    add_section_heading(doc, "4.4. Hướng phát triển")
    add_paragraph(doc, "Trong giai đoạn tiếp theo, hệ thống có thể được phát triển theo hướng triển khai production, hoàn thiện upload Google Drive, đồng bộ logic khuyến mãi giữa frontend và backend, bổ sung cổng PayOS nếu cần, xây dựng phân quyền chi tiết hơn cho photographer và bổ sung báo cáo doanh thu theo tháng/quý.")
    add_paragraph(doc, "Ngoài ra, hệ thống nên có bộ kiểm thử tự động cho các nghiệp vụ quan trọng. Với thanh toán, cần test các trường hợp chữ ký sai, giao dịch thất bại, booking hết hạn và thanh toán lặp. Với admin, cần test quy tắc chuyển trạng thái để tránh sai lệch dữ liệu vận hành.")

    doc.add_page_break()
    add_part(doc, "PHẦN 3. KẾT LUẬN")
    add_section_heading(doc, "3.1. Kết quả đạt được")
    add_paragraph(doc, "Đề tài đã xây dựng được website quản lý và đặt lịch dịch vụ chụp ảnh cho Cao Hiển Studio với đầy đủ các module trọng tâm. Khách hàng có thể tiếp cận thông tin dịch vụ, xem album, đăng ký tài khoản, đặt lịch và thanh toán cọc. Quản trị viên có thể theo dõi dashboard, quản lý booking, dịch vụ, album, danh mục động và khách hàng.")
    add_paragraph(doc, "Về kỹ thuật, hệ thống áp dụng kiến trúc tách biệt frontend React/Vite và backend Node.js/Express, sử dụng MongoDB để lưu trữ dữ liệu, JWT để xác thực, VNPay để thanh toán, Nodemailer để gửi OTP, Google Drive để hiển thị album và Gemini AI để hỗ trợ tư vấn khách hàng.")

    add_section_heading(doc, "3.2. Kiến thức và kinh nghiệm thu được")
    add_paragraph(doc, "Thông qua đề tài, sinh viên rèn luyện được kỹ năng phân tích nghiệp vụ, thiết kế dữ liệu, xây dựng API, tổ chức giao diện frontend, xử lý xác thực, tích hợp thanh toán, tổ chức danh mục động, tự động hóa email và kiểm thử luồng chức năng. Đề tài cũng giúp hiểu rõ hơn cách chuyển một bài toán vận hành thực tế của studio thành các module phần mềm có thể triển khai.")

    add_section_heading(doc, "3.3. Kết luận chung")
    add_paragraph(doc, "Website Cao Hiển Studio là một sản phẩm có tính ứng dụng thực tiễn, phù hợp với nhu cầu số hóa quy trình tư vấn, đặt lịch và quản lý dịch vụ chụp ảnh. Mặc dù vẫn còn một số hạn chế cần hoàn thiện, hệ thống đã tạo nền tảng tốt để mở rộng thành sản phẩm triển khai thật cho studio trong tương lai.")

    doc.add_page_break()
    add_center(doc, "TÀI LIỆU THAM KHẢO", size=16, bold=True, after=12)
    references = [
        "[1] React Documentation, https://react.dev/.",
        "[2] Vite Documentation, https://vitejs.dev/.",
        "[3] Express.js Documentation, https://expressjs.com/.",
        "[4] MongoDB Documentation, https://www.mongodb.com/docs/.",
        "[5] Mongoose Documentation, https://mongoosejs.com/docs/.",
        "[6] JSON Web Token Introduction, https://jwt.io/introduction.",
        "[7] Ant Design Documentation, https://ant.design/docs/react/introduce.",
        "[8] VNPay Payment Gateway Documentation, https://sandbox.vnpayment.vn/apis/.",
        "[9] Google Drive API Documentation, https://developers.google.com/drive/api.",
        "[10] Gemini API Documentation, https://ai.google.dev/docs.",
        "[11] Nodemailer Documentation, https://nodemailer.com/about/.",
        "[12] Open-Meteo API Documentation, https://open-meteo.com/en/docs.",
        "[13] Photon Geocoder, https://photon.komoot.io/.",
    ]
    for ref in references:
        p = add_paragraph(doc, ref, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)


def build():
    ensure_dirs()
    logo = create_logo()
    images = {
        "architecture": create_architecture_diagram(),
        "usecase": create_usecase_diagram(),
        "booking_flow": create_booking_flow_diagram(),
        "erd": create_erd_diagram(),
    }

    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0], border=True)
    add_front_matter(doc, logo)
    add_body(doc, images)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
